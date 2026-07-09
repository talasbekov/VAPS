"""Story 6.3 — контракт данных и чистый рендерер официального .docx расхода.

Канон вёрстки — addendum §8 (+ FR-17): секция A4 landscape, заголовок
``{Подразделение} ЖЕКЕ ҚҰРАМЫНЫҢ САПТЫҚ ТІЗІМІ {ДД.ММ.ГГГГ} ЖЫЛҒЫ`` (16pt,
жирный — Д10), шапка и числа таблицы 12pt, списки членов в ячейках 8pt,
строка «ИТОГО» жирным; ``Times New Roman`` проставлен ЯВНО на каждом run
(дефолтный Calibri шаблона python-docx не допускается — Ловушка №10).

Рендерер — ЧИСТАЯ функция (зеркало ``strength_render``): без ORM, без
wall-clock, без сети и записи на диск; результат — ``bytes``. ``python-docx``
импортируется ЛЕНИВО внутри ``generate_expense_docx`` — импорт модуля живёт
без установленного пакета. Рендерер НЕ считает и НЕ пересортировывает: числа
и порядок строк приходят готовыми от билдера (данные = derive(снапшот),
Story 6.3 Task 3), формулы сходимости держит derive и рантайм-ассерт выпуска
(6.5) — дублировать их здесь запрещено (Д8). Усечение длинных списков членов
(cap ``CELL_MAX_MEMBERS`` + строка ``… ещё N``) — единственная «логика» рендера.

Датаклассы контракта живут ЗДЕСЬ (не в operations): AST-гвард изоляции банит
``apps.operations.*`` в app documents, а стрелка «documents ← operations»
разрешена — операционный билдер импортирует ЭТИ типы (Ловушка №3).
"""

import io
from dataclasses import dataclass
from datetime import date

# Порядок ключей derive-колонок: канон addendum §8 + хвост Д3
# (BEFORE_DUTY/OTHER/PENDING обязаны быть видимы — иначе Σ видимых колонок
# ≠ «По списку» и формула ломается В ДОКУМЕНТЕ). Дубль ключей с
# operations.REPORT_COLUMNS — ОСОЗНАННЫЙ (границу не пересекаем); дрейф ловит
# sync-тест в apps/operations/submissions/tests (там обе стороны видны легально).
DOCX_COLUMNS = (
    "IN_SERVICE",
    "ON_DUTY",
    "AFTER_DUTY",
    "COMMAND",
    "TRAINING",
    "VACATION",
    "SICK",
    "ATTACHED",
    "DETACHED",
    "BEFORE_DUTY",
    "OTHER",
    "PENDING",
)

# Русские лейблы шапки: canonical source — addendum §8 (литеральная копия),
# хвост «Перед дежурством»/«Иное» — выжимка ТЗ §77.3, «Уточняется» — стори 3.9.
DOCX_COLUMN_LABELS = {
    "IN_SERVICE": "В строю",
    "ON_DUTY": "На дежурстве",
    "AFTER_DUTY": "После дежурства",
    "COMMAND": "В командировке",
    "TRAINING": "Учёба/соревнования/конференция",
    "VACATION": "В отпуске",
    "SICK": "На больничном",
    "ATTACHED": "Прикомандирован",
    "DETACHED": "Откомандирован",
    "BEFORE_DUTY": "Перед дежурством",
    "OTHER": "Иное",
    "PENDING": "Уточняется",
}

# Cap строк-членов в ячейке (выжимка ТЗ §8.5, DAILY_REPORT_CELL_MAX_NAMES).
CELL_MAX_MEMBERS = 20

FONT_NAME = "Times New Roman"
TITLE_SIZE_PT = 16
TABLE_SIZE_PT = 12
MEMBER_SIZE_PT = 8

_FIXED_HEAD = ("№", "Управление", "По штату", "По списку", "Вакансии")
_TOTALS_LABEL = "ИТОГО"  # заглавными — вид донора utils.py (Д6)
_DATE_FORMAT = "%d.%m.%Y"


@dataclass(frozen=True)
class ExpenseCellMember:
    """Строка члена ячейки: ``{rank} {full_name} — {период winner-факта}``."""

    rank: str
    full_name: str
    date_start: date
    date_end: date


@dataclass(frozen=True)
class ExpenseCell:
    """Ячейка статусной колонки: число + ПОЛНЫЙ список членов (усечение до
    ``CELL_MAX_MEMBERS`` — зона рендера, count всегда согласован с полным
    списком — cross-assert билдера)."""

    count: int
    members: tuple


@dataclass(frozen=True)
class ExpenseRow:
    """Одна строка документа. ``cells`` — ключ DOCX_COLUMNS → ExpenseCell для
    всех колонок, КРОМЕ ATTACHED: прикомандированные вне «По списку», потому
    живут отдельным полем ``attached`` (рендер «+N»)."""

    name: str
    staff_total: int
    list_total: int
    vacancies: int
    cells: dict
    attached: ExpenseCell


@dataclass(frozen=True)
class ExpenseTotals:
    """Строка «ИТОГО» — только числа, БЕЗ списков членов."""

    staff_total: int
    list_total: int
    vacancies: int
    columns: dict
    attached: int


@dataclass(frozen=True)
class ExpenseDocumentData:
    division_title: str
    business_date: date
    rows: list
    totals: ExpenseTotals


def _member_line(member):
    """``{rank} {full_name} — {ДД.ММ.ГГГГ}–{ДД.ММ.ГГГГ}``; пустой rank — без
    ведущего пробела."""
    name = " ".join(filter(None, [member.rank, member.full_name]))
    return (
        f"{name} — {member.date_start.strftime(_DATE_FORMAT)}"
        f"–{member.date_end.strftime(_DATE_FORMAT)}"
    )


def generate_expense_docx(data):
    """``ExpenseDocumentData`` → байты официального .docx (канон §8).

    Ленивая загрузка python-docx — зеркало ``strength_render.build_workbook``.
    """
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt

    def styled(paragraph, text, size_pt, bold=False):
        """Run с ЯВНЫМИ font.name/size (Ловушка №10): дефолт шаблона —
        Calibri; ``font.name`` кроет ascii/hAnsi, ``w:cs`` добиваем через
        rFonts для полноты покрытия кириллицы/казахских глифов."""
        run = paragraph.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(size_pt)
        if bold:
            run.bold = True
        rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        rfonts.set(qn("w:cs"), FONT_NAME)
        return run

    def fill_status_cell(cell, key, cell_data, bold=False):
        """Число (ATTACHED — «+N») 12pt + до CELL_MAX_MEMBERS строк членов 8pt
        отдельными параграфами; хвост усечения — ``… ещё N``."""
        count_text = (
            f"+{cell_data.count}" if key == "ATTACHED" else str(cell_data.count)
        )
        styled(cell.paragraphs[0], count_text, TABLE_SIZE_PT, bold=bold)
        shown = cell_data.members[:CELL_MAX_MEMBERS]
        for member in shown:
            styled(cell.add_paragraph(), _member_line(member), MEMBER_SIZE_PT)
        hidden = len(cell_data.members) - len(shown)
        if hidden > 0:
            styled(cell.add_paragraph(), f"… ещё {hidden}", MEMBER_SIZE_PT)

    document = Document()

    # A4 landscape: дефолтный шаблон — Letter portrait; библиотека при смене
    # orientation размеры сама НЕ свопает (проверено эмпирически на 1.2.0) —
    # ставим и ориентацию, и габариты A4 явно.
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)

    title = document.add_paragraph()
    styled(
        title,
        f"{data.division_title} ЖЕКЕ ҚҰРАМЫНЫҢ САПТЫҚ ТІЗІМІ "
        f"{data.business_date.strftime(_DATE_FORMAT)} ЖЫЛҒЫ",
        TITLE_SIZE_PT,
        bold=True,
    )

    header = _FIXED_HEAD + tuple(DOCX_COLUMN_LABELS[key] for key in DOCX_COLUMNS)
    table = document.add_table(rows=1, cols=len(header))
    for cell, label in zip(table.rows[0].cells, header):
        styled(cell.paragraphs[0], label, TABLE_SIZE_PT, bold=True)

    for number, row in enumerate(data.rows, start=1):
        cells = table.add_row().cells
        styled(cells[0].paragraphs[0], str(number), TABLE_SIZE_PT)
        styled(cells[1].paragraphs[0], row.name, TABLE_SIZE_PT)
        styled(cells[2].paragraphs[0], str(row.staff_total), TABLE_SIZE_PT)
        styled(cells[3].paragraphs[0], str(row.list_total), TABLE_SIZE_PT)
        styled(cells[4].paragraphs[0], str(row.vacancies), TABLE_SIZE_PT)
        for offset, key in enumerate(DOCX_COLUMNS, start=len(_FIXED_HEAD)):
            cell_data = row.attached if key == "ATTACHED" else row.cells[key]
            fill_status_cell(cells[offset], key, cell_data)

    totals = data.totals
    totals_cells = table.add_row().cells
    styled(totals_cells[1].paragraphs[0], _TOTALS_LABEL, TABLE_SIZE_PT, bold=True)
    styled(
        totals_cells[2].paragraphs[0], str(totals.staff_total), TABLE_SIZE_PT, bold=True
    )
    styled(
        totals_cells[3].paragraphs[0], str(totals.list_total), TABLE_SIZE_PT, bold=True
    )
    styled(
        totals_cells[4].paragraphs[0], str(totals.vacancies), TABLE_SIZE_PT, bold=True
    )
    for offset, key in enumerate(DOCX_COLUMNS, start=len(_FIXED_HEAD)):
        if key == "ATTACHED":
            text = f"+{totals.attached}"
        else:
            text = str(totals.columns[key])
        styled(totals_cells[offset].paragraphs[0], text, TABLE_SIZE_PT, bold=True)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
