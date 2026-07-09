"""Тесты чистого рендерера .docx расхода (Story 6.3, AC-1/2/4/5).

БЕЗ ``django_db``: рендерер — чистая функция (datacls-вход → bytes), читаем
сгенерированное обратно через ``docx.Document(BytesIO(result))``. Числа входа —
литералы: рендерер формулы сходимости НЕ проверяет (Д8 — это зона derive/6.5),
поэтому вход не обязан сходиться.
"""

from datetime import date
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT

from apps.documents.generators import (
    CELL_MAX_MEMBERS,
    DOCX_COLUMN_LABELS,
    DOCX_COLUMNS,
    FONT_NAME,
    ExpenseCell,
    ExpenseCellMember,
    ExpenseDocumentData,
    ExpenseRow,
    ExpenseTotals,
    generate_expense_docx,
)

D = date(2026, 7, 8)

# Индексы фиксированной шапки: № | Управление | По штату | По списку | Вакансии
# | 12 колонок DOCX_COLUMNS. Смещение статусной колонки = 5 + позиция ключа.
_FIXED_HEAD = ("№", "Управление", "По штату", "По списку", "Вакансии")


def _col(key):
    return len(_FIXED_HEAD) + DOCX_COLUMNS.index(key)


def _member(rank="капитан", full_name="Иванов Иван", start=None, end=None):
    return ExpenseCellMember(
        rank=rank,
        full_name=full_name,
        date_start=start or date(2026, 7, 1),
        date_end=end or date(2026, 7, 10),
    )


def _cells(**counts):
    """Все 11 статусных ячеек (без ATTACHED); именованные — с содержимым."""
    cells = {}
    for key in DOCX_COLUMNS:
        if key == "ATTACHED":
            continue
        cells[key] = counts.get(key, ExpenseCell(count=0, members=()))
    return cells


def _totals(**overrides):
    columns = {key: 0 for key in DOCX_COLUMNS if key != "ATTACHED"}
    columns.update(overrides.pop("columns", {}))
    defaults = {"staff_total": 10, "list_total": 8, "vacancies": 2, "attached": 1}
    defaults.update(overrides)
    return ExpenseTotals(columns=columns, **defaults)


def _data(rows=None, totals=None, division_title="1-е управление"):
    if rows is None:
        rows = [
            ExpenseRow(
                name="1-е управление",
                staff_total=10,
                list_total=8,
                vacancies=2,
                cells=_cells(
                    IN_SERVICE=ExpenseCell(count=4, members=()),
                    ON_DUTY=ExpenseCell(count=1, members=(_member(),)),
                    TRAINING=ExpenseCell(
                        count=2,
                        members=(
                            _member(full_name="Петров Пётр"),
                            _member(rank="майор", full_name="Сидоров Сидор"),
                        ),
                    ),
                    VACATION=ExpenseCell(count=1, members=(_member(rank=""),)),
                ),
                attached=ExpenseCell(count=1, members=(_member(),)),
            )
        ]
    if totals is None:
        totals = _totals(
            columns={"IN_SERVICE": 4, "ON_DUTY": 1, "TRAINING": 2, "VACATION": 1}
        )
    return ExpenseDocumentData(
        division_title=division_title, business_date=D, rows=rows, totals=totals
    )


def _generated(data):
    result = generate_expense_docx(data)
    assert isinstance(result, bytes)
    return Document(BytesIO(result))


def _all_runs(document):
    runs = [run for paragraph in document.paragraphs for run in paragraph.runs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    runs.extend(paragraph.runs)
    return runs


def test_bytes_roundtrip_smoke():
    # bytes-результат — валидный zip/docx: открывается python-docx без ошибок.
    result = generate_expense_docx(_data())
    assert isinstance(result, bytes)
    assert result[:2] == b"PK"
    Document(BytesIO(result))  # не бросает


def test_section_is_a4_landscape():
    section = _generated(_data()).sections[0]
    assert section.orientation == WD_ORIENT.LANDSCAPE
    # Габариты в OOXML хранятся в twips — точные EMU не round-trip'ятся,
    # сверяем миллиметры с округлением.
    assert round(section.page_width.mm) == 297
    assert round(section.page_height.mm) == 210
    assert section.page_width > section.page_height


def test_title_text_and_style():
    document = _generated(_data())
    title = document.paragraphs[0]
    assert title.text == (
        "1-е управление ЖЕКЕ ҚҰРАМЫНЫҢ САПТЫҚ ТІЗІМІ 08.07.2026 ЖЫЛҒЫ"
    )
    assert title.runs, "заголовок без runs"
    for run in title.runs:
        assert run.font.size.pt == 16
        assert run.bold is True
        assert run.font.name == FONT_NAME


def test_header_row_full_label_order():
    table = _generated(_data()).tables[0]
    labels = [cell.text for cell in table.rows[0].cells]
    assert labels == list(_FIXED_HEAD) + [
        DOCX_COLUMN_LABELS[key] for key in DOCX_COLUMNS
    ]
    # Литеральная сверка канона AC-2 (§8-порядок + хвост Д3) — от дрейфа констант.
    assert labels == [
        "№",
        "Управление",
        "По штату",
        "По списку",
        "Вакансии",
        "В строю",
        "На дежурстве",
        "После дежурства",
        "В командировке",
        "Учёба/соревнования/конференция",
        "В отпуске",
        "На больничном",
        "Прикомандирован",
        "Откомандирован",
        "Перед дежурством",
        "Иное",
        "Уточняется",
    ]


def test_body_row_values_from_input():
    table = _generated(_data()).tables[0]
    row = table.rows[1]
    assert row.cells[0].text == "1"
    assert row.cells[1].text == "1-е управление"
    assert row.cells[2].text == "10"
    assert row.cells[3].text == "8"
    assert row.cells[4].text == "2"
    assert row.cells[_col("IN_SERVICE")].text == "4"
    # Нулевая колонка — просто "0", без списков.
    assert row.cells[_col("SICK")].text == "0"


def test_status_cell_count_plus_member_lines():
    table = _generated(_data()).tables[0]
    cell = table.rows[1].cells[_col("TRAINING")]
    texts = [paragraph.text for paragraph in cell.paragraphs]
    assert texts == [
        "2",
        "капитан Петров Пётр — 01.07.2026–10.07.2026",
        "майор Сидоров Сидор — 01.07.2026–10.07.2026",
    ]
    count_runs = cell.paragraphs[0].runs
    member_runs = [run for p in cell.paragraphs[1:] for run in p.runs]
    assert all(run.font.size.pt == 12 for run in count_runs)
    assert member_runs and all(run.font.size.pt == 8 for run in member_runs)


def test_member_line_empty_rank_has_no_leading_space():
    table = _generated(_data()).tables[0]
    cell = table.rows[1].cells[_col("VACATION")]
    assert cell.paragraphs[1].text == "Иванов Иван — 01.07.2026–10.07.2026"


def test_attached_rendered_as_plus_n_with_members():
    table = _generated(_data()).tables[0]
    cell = table.rows[1].cells[_col("ATTACHED")]
    assert cell.paragraphs[0].text == "+1"
    assert cell.paragraphs[1].text == "капитан Иванов Иван — 01.07.2026–10.07.2026"


def test_empty_members_render_only_number():
    table = _generated(_data()).tables[0]
    cell = table.rows[1].cells[_col("IN_SERVICE")]
    assert len(cell.paragraphs) == 1
    assert cell.text == "4"


def test_truncation_over_cap_members():
    members = tuple(
        _member(full_name=f"Сотрудник {n:02d}") for n in range(CELL_MAX_MEMBERS + 1)
    )
    row = ExpenseRow(
        name="Отдел",
        staff_total=30,
        list_total=25,
        vacancies=5,
        cells=_cells(ON_DUTY=ExpenseCell(count=len(members), members=members)),
        attached=ExpenseCell(count=0, members=()),
    )
    table = _generated(_data(rows=[row], totals=_totals())).tables[0]
    cell = table.rows[1].cells[_col("ON_DUTY")]
    texts = [paragraph.text for paragraph in cell.paragraphs]
    assert texts[0] == str(CELL_MAX_MEMBERS + 1)
    assert len(texts) == 1 + CELL_MAX_MEMBERS + 1  # count + 20 членов + хвост
    assert "Сотрудник 00" in texts[1]
    assert texts[-1] == "… ещё 1"


def test_totals_row_values_and_bold():
    table = _generated(_data()).tables[0]
    totals_row = table.rows[-1]
    assert totals_row.cells[1].text == "ИТОГО"
    assert totals_row.cells[2].text == "10"
    assert totals_row.cells[3].text == "8"
    assert totals_row.cells[4].text == "2"
    assert totals_row.cells[_col("IN_SERVICE")].text == "4"
    assert totals_row.cells[_col("ATTACHED")].text == "+1"
    runs = [
        run
        for cell in totals_row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
    ]
    assert runs, "строка ИТОГО без runs"
    assert all(run.bold is True for run in runs)
    assert all(run.font.size.pt == 12 for run in runs)


def test_rows_rendered_in_given_order_without_resort():
    rows = [
        ExpenseRow(
            name=name,
            staff_total=1,
            list_total=1,
            vacancies=0,
            cells=_cells(IN_SERVICE=ExpenseCell(count=1, members=())),
            attached=ExpenseCell(count=0, members=()),
        )
        for name in ("Яблоко", "Абрикос")  # нарочно НЕ по алфавиту
    ]
    table = _generated(_data(rows=rows, totals=_totals())).tables[0]
    assert [table.rows[1].cells[0].text, table.rows[1].cells[1].text] == [
        "1",
        "Яблоко",
    ]
    assert [table.rows[2].cells[0].text, table.rows[2].cells[1].text] == [
        "2",
        "Абрикос",
    ]


def test_every_run_carries_explicit_font_and_canonical_size():
    document = _generated(_data())
    runs = _all_runs(document)
    assert runs
    for run in runs:
        assert run.font.name == FONT_NAME  # не дефолтный Calibri (AC-1)
        assert run.font.size is not None
        assert run.font.size.pt in (16, 12, 8)
