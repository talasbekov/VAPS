"""QA-судьи чистого рендерера .docx расхода (Story 6.3, bmad-qa-generate-e2e-tests).

Дополняют dev-сьют по мутационному принципу («какая правка кода не краснит
НИ ОДИН тест»): w:cs-шрифт казахских глифов (Ловушка №10 — dev судит только
``font.name``), граница cap РОВНО в ``CELL_MAX_MEMBERS``, стиль хвоста
усечения, row-агностичность рендерера (multi-row потребитель — свод 6.5)
и пин видимого «+0» пустой прикомандированности. БЕЗ ``django_db`` (чистая
функция); dev-модуль ``test_expense_docx_generator.py`` не тронут.
"""

from datetime import date
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn

from apps.documents.generators import (
    CELL_MAX_MEMBERS,
    DOCX_COLUMNS,
    FONT_NAME,
    ExpenseCell,
    ExpenseCellMember,
    ExpenseDocumentData,
    ExpenseRow,
    ExpenseTotals,
    generate_expense_docx,
)

D = date(2026, 7, 8)

# Смещение статусной колонки: № | Управление | По штату | По списку | Вакансии.
_FIXED_HEAD_LEN = 5


def _col(key):
    return _FIXED_HEAD_LEN + DOCX_COLUMNS.index(key)


def _member(full_name):
    return ExpenseCellMember(
        rank="капитан",
        full_name=full_name,
        date_start=date(2026, 7, 1),
        date_end=date(2026, 7, 10),
    )


def _row(members=(), attached=ExpenseCell(count=0, members=())):
    cells = {
        key: ExpenseCell(count=0, members=())
        for key in DOCX_COLUMNS
        if key != "ATTACHED"
    }
    if members:
        cells["ON_DUTY"] = ExpenseCell(count=len(members), members=tuple(members))
    return ExpenseRow(
        name="Отдел К",
        staff_total=30,
        list_total=25,
        vacancies=5,
        cells=cells,
        attached=attached,
    )


def _data(rows, attached_total=0):
    return ExpenseDocumentData(
        division_title="Отдел К",
        business_date=D,
        rows=rows,
        totals=ExpenseTotals(
            staff_total=30,
            list_total=25,
            vacancies=5,
            columns={key: 0 for key in DOCX_COLUMNS if key != "ATTACHED"},
            attached=attached_total,
        ),
    )


def _document(data):
    result = generate_expense_docx(data)
    assert isinstance(result, bytes)
    return Document(BytesIO(result))


def test_exactly_cap_members_render_all_without_tail():
    # Граница cap: РОВНО CELL_MAX_MEMBERS членов — все видимы, хвоста
    # «… ещё N» НЕТ. Dev судит только 21 → 20 + хвост; rewrite-мутация
    # класса «if len(members) >= CAP: хвост» на границе была невидима.
    members = tuple(_member(f"Сотрудник {n:02d}") for n in range(CELL_MAX_MEMBERS))
    document = _document(_data([_row(members=members)]))
    cell = document.tables[0].rows[1].cells[_col("ON_DUTY")]
    texts = [paragraph.text for paragraph in cell.paragraphs]
    assert texts[0] == str(CELL_MAX_MEMBERS)
    assert len(texts) == 1 + CELL_MAX_MEMBERS
    assert not any(text.startswith("… ещё") for text in texts)
    assert "Сотрудник 19" in texts[-1]


def test_overcap_member_and_tail_runs_styled_tnr_8pt():
    # Хвост «… ещё N» и строки членов УСЕЧЁННОЙ ячейки — 8pt/TNR: dev-судья
    # шрифтов бежит на данных без усечения, нестилизованный хвост
    # (голый ``add_paragraph(text)`` вместо styled) не краснил ничего.
    members = tuple(_member(f"Сотрудник {n:02d}") for n in range(CELL_MAX_MEMBERS + 3))
    document = _document(_data([_row(members=members)]))
    cell = document.tables[0].rows[1].cells[_col("ON_DUTY")]
    assert cell.paragraphs[-1].text == "… ещё 3"
    member_runs = [run for p in cell.paragraphs[1:] for run in p.runs]
    assert member_runs
    for run in member_runs:
        assert run.font.name == FONT_NAME
        assert run.font.size.pt == 8


def test_every_run_sets_cs_font_for_kazakh_glyphs():
    # Ловушка №10: ``font.name`` кроет только ascii/hAnsi; полноту покрытия
    # казахских глифов (Қ/Ұ/Ә/Ғ) держит атрибут w:cs в rFonts. Dev судит
    # только font.name — мутация «удалить rfonts.set(w:cs)» была невидима.
    document = _document(
        _data(
            [
                _row(
                    members=(_member("Иванов Иван"),),
                    attached=ExpenseCell(count=1, members=(_member("Гость"),)),
                )
            ],
            attached_total=1,
        )
    )
    runs = [run for paragraph in document.paragraphs for run in paragraph.runs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    runs.extend(paragraph.runs)
    assert runs
    for run in runs:
        rpr = run._element.rPr
        assert rpr is not None, run.text
        assert rpr.rFonts is not None, run.text
        assert rpr.rFonts.get(qn("w:cs")) == FONT_NAME, run.text


def test_renderer_row_agnostic_zero_and_two_rows_totals_last():
    # Рендерер row-агностичен (однострочность Д11 — контракт БИЛДЕРА, свод
    # 6.5 подаст больше строк): 0 строк — шапка + ИТОГО без падения; 2
    # строки — «ИТОГО» ровно один раз и ПОСЛЕДНЕЙ строкой таблицы.
    empty = _document(_data([])).tables[0]
    assert len(empty.rows) == 2
    assert empty.rows[-1].cells[1].text == "ИТОГО"

    two = _document(_data([_row(), _row()])).tables[0]
    assert len(two.rows) == 4
    labels = [row.cells[1].text for row in two.rows]
    assert labels.count("ИТОГО") == 1
    assert labels[-1] == "ИТОГО"


def test_zero_attached_renders_plus_zero_pin():
    # Пин видимого поведения: нулевая прикомандированность — «+0» (формат
    # «+N» не имеет особого нуля), тело и ИТОГО согласованы. Если заказчик
    # захочет «0»/пусто — менять ОСОЗНАННО вместе с этим судьёй.
    document = _document(_data([_row()]))
    table = document.tables[0]
    body_cell = table.rows[1].cells[_col("ATTACHED")]
    assert body_cell.text == "+0"
    assert len(body_cell.paragraphs) == 1
    assert table.rows[-1].cells[_col("ATTACHED")].text == "+0"
