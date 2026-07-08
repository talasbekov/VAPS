"""Тест эквивалентности чисел четырёх форматов расхода (Story 6.4, AC-1).

ОДНА фикстура ≥2 строк → ``generate_expense_docx/xlsx/csv/pdf`` → каждое
числовое значение (Штат/Список/Вакансии/12 колонок/строка ИТОГО; ATTACHED
нормализуется ``+N`` → N) идентично во всех четырёх форматах поле-в-поле,
порядок строк одинаков (рендереры НЕ пересортировывают).

Все числа фикстуры — попарно различные трёхзначные из НЕПЕРЕСЕКАЮЩЕГОСЯ
диапазона (не встречаются во фрагментах дат периодов членов 01/07/10/2026,
№ строк 1/2 и дате титула 08.07.2026) — для PDF присутствие числа-токена
означает его маппинг, а не коллизию (анти-C2). docx≡xlsx≡csv сравниваются
позиционно; pdf ⊇ те же числа отдельными токенами.

БЕЗ ``django_db``: чистые функции, вход — литеральный контракт (рендереры
формулы НЕ проверяют — числа не обязаны сходиться).
"""

import csv
import itertools
import re
from datetime import date
from io import BytesIO, StringIO

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from apps.documents.generators import (
    DOCX_COLUMNS,
    ExpenseCell,
    ExpenseCellMember,
    ExpenseDocumentData,
    ExpenseRow,
    ExpenseTotals,
    generate_expense_csv,
    generate_expense_docx,
    generate_expense_pdf,
    generate_expense_xlsx,
)

D = date(2026, 7, 8)
_FIXED_HEAD_LEN = 5  # № | Управление | По штату | По списку | Вакансии


def _member(rank="капитан", full_name="Иванов Иван"):
    return ExpenseCellMember(
        rank=rank,
        full_name=full_name,
        date_start=date(2026, 7, 1),
        date_end=date(2026, 7, 10),
    )


def _fixture():
    """≥2 строк; каждое число — уникальное трёхзначное (счётчик с 101);
    включает members-ячейку (ON_DUTY строки 1, ATTACHED строки 1), пустую
    ячейку (PENDING строки 2 — count 0 без членов) и ATTACHED>0 в обеих
    строках. Возвращает (data, used) — used БЕЗ нуля."""
    sequence = itertools.count(101)
    used = []

    def n():
        value = next(sequence)
        used.append(value)
        return value

    def cells(members_key=None, zero_key=None):
        result = {}
        for key in DOCX_COLUMNS:
            if key == "ATTACHED":
                continue
            if key == zero_key:
                result[key] = ExpenseCell(count=0, members=())
            elif key == members_key:
                result[key] = ExpenseCell(
                    count=n(),
                    members=(
                        _member(full_name="Петров Пётр"),
                        _member(rank="майор", full_name="Сидоров Сидор"),
                    ),
                )
            else:
                result[key] = ExpenseCell(count=n(), members=())
        return result

    rows = [
        ExpenseRow(
            name="Альфа",
            staff_total=n(),
            list_total=n(),
            vacancies=n(),
            cells=cells(members_key="ON_DUTY"),
            attached=ExpenseCell(count=n(), members=(_member(),)),
        ),
        ExpenseRow(
            name="Бета",
            staff_total=n(),
            list_total=n(),
            vacancies=n(),
            cells=cells(zero_key="PENDING"),
            attached=ExpenseCell(count=n(), members=()),
        ),
    ]
    totals = ExpenseTotals(
        staff_total=n(),
        list_total=n(),
        vacancies=n(),
        columns={key: n() for key in DOCX_COLUMNS if key != "ATTACHED"},
        attached=n(),
    )
    data = ExpenseDocumentData(
        division_title="Управление Эталон", business_date=D, rows=rows, totals=totals
    )
    return data, used


def _int(text):
    """Нормализация ATTACHED: ``+N`` → N."""
    return int(str(text).lstrip("+"))


def _matrix_docx(raw):
    """Числа документа по строкам: [Штат, Список, Вакансии, 12 колонок].
    Членская ячейка несёт даты членов в ``cell.text`` — берём ПЕРВЫЙ
    параграф ячейки (count)."""
    table = Document(BytesIO(raw)).tables[0]
    matrix = []
    for row in table.rows[1:]:
        numbers = [_int(row.cells[column].text) for column in (2, 3, 4)]
        for offset in range(len(DOCX_COLUMNS)):
            cell = row.cells[_FIXED_HEAD_LEN + offset]
            numbers.append(_int(cell.paragraphs[0].text))
        matrix.append(numbers)
    return matrix


def _matrix_xlsx(raw):
    """Текст-ячейки (члены/ATTACHED) — первая строка текста; безчленные —
    уже int."""
    sheet = load_workbook(BytesIO(raw)).active
    matrix = []
    for sheet_row in range(3, sheet.max_row + 1):
        numbers = []
        for column in (3, 4, 5):
            value = sheet.cell(row=sheet_row, column=column).value
            numbers.append(_int(value) if value is not None else None)
        for offset in range(len(DOCX_COLUMNS)):
            value = sheet.cell(row=sheet_row, column=6 + offset).value
            if isinstance(value, str):
                value = value.splitlines()[0]
            numbers.append(_int(value))
        matrix.append(numbers)
    return matrix


def _matrix_csv(raw):
    rows = list(csv.reader(StringIO(raw.decode("utf-8-sig")), delimiter=";"))
    matrix = []
    for record in rows[2:]:
        matrix.append([_int(field) for field in record[2:]])
    return matrix


def test_numbers_identical_across_docx_xlsx_csv_positionally():
    data, _ = _fixture()
    docx_matrix = _matrix_docx(generate_expense_docx(data))
    xlsx_matrix = _matrix_xlsx(generate_expense_xlsx(data))
    csv_matrix = _matrix_csv(generate_expense_csv(data))
    # Поле-в-поле, позиционно, включая порядок строк и строку ИТОГО.
    assert docx_matrix == xlsx_matrix == csv_matrix
    assert len(docx_matrix) == 3  # 2 строки данных + ИТОГО


def test_matrices_mirror_the_input_contract():
    # Матрица — не только межформатная: она равна самому контракту
    # (рендереры не искажают и не пересортировывают вход).
    data, _ = _fixture()
    expected = []
    for row in data.rows:
        numbers = [row.staff_total, row.list_total, row.vacancies]
        for key in DOCX_COLUMNS:
            cell = row.attached if key == "ATTACHED" else row.cells[key]
            numbers.append(cell.count)
        expected.append(numbers)
    totals = data.totals
    totals_numbers = [totals.staff_total, totals.list_total, totals.vacancies]
    for key in DOCX_COLUMNS:
        totals_numbers.append(
            totals.attached if key == "ATTACHED" else totals.columns[key]
        )
    expected.append(totals_numbers)
    assert _matrix_docx(generate_expense_docx(data)) == expected


def test_pdf_contains_every_unique_number_as_token():
    data, used = _fixture()
    reader = PdfReader(BytesIO(generate_expense_pdf(data)))
    text = "\n".join(page.extract_text() for page in reader.pages)
    tokens = set(re.split(r"\D+", text))
    # Числа уникальны и вне служебных диапазонов — presence == mapping.
    missing = [number for number in used if str(number) not in tokens]
    assert missing == []
    # Нормализация +N → N: ATTACHED-значения присутствуют и с плюсом.
    normalized = "".join(text.split())
    assert f"+{data.rows[0].attached.count}" in normalized
    assert f"+{data.totals.attached}" in normalized


def test_fixture_numbers_are_pairwise_distinct_three_digit():
    # Самопроверка фикстуры (анти-C2): уникальность и непересечение со
    # служебными токенами (даты членов, № строк, дата титула).
    _, used = _fixture()
    assert len(used) == len(set(used))
    assert all(100 <= number <= 999 for number in used)
    reserved = {1, 2, 7, 8, 10, 2026}
    assert not reserved & set(used)
