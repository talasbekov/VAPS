"""E2e четырёх форматов расхода от реального снапшота (Story 6.4, AC-1).

Снапшот ТОЛЬКО ``build_division_snapshot`` (урок ревью 5.4b, руками не
лепим) → ``build_expense_document`` → все четыре генератора → числа каждого
формата равны полям НЕЗАВИСИМОГО ``derive_report`` от тех же входов.
Данные — зеркало e2e 6.3 (фикстуры — копия ``test_expense_document.py``):
multi-код колонка (STUDY + CONFERENCE → TRAINING), DUTY, VACATION, ATTACHED
и пара сотрудников без факта (В строю).

``django_db`` обычный, БЕЗ ``transaction=True`` (teardown-ловушка 6.2).
"""

import csv
import itertools
import re
from datetime import date
from io import BytesIO, StringIO

import pytest
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from apps.core.models import (
    Division,
    DivisionType,
    Employee,
    Organization,
    Rank,
)
from apps.documents.generators import (
    DOCX_COLUMNS,
    generate_expense_csv,
    generate_expense_docx,
    generate_expense_pdf,
    generate_expense_xlsx,
)
from apps.operations.statuses.models import EmployeeStatus
from apps.operations.statuses.services.strength_report import derive_report
from apps.operations.submissions.expense_document import build_expense_document
from apps.operations.submissions.services.snapshot import build_division_snapshot

pytestmark = pytest.mark.django_db

D = date(2026, 7, 8)
_iin_seq = itertools.count(900)

# Смещение статусной колонки: № | Управление | По штату | По списку |
# Вакансии | 12 колонок DOCX_COLUMNS.
_FIXED_HEAD_LEN = 5


@pytest.fixture
def org():
    return Organization.objects.create(name="Орг", code="ORG-EXPF")


@pytest.fixture
def division(org):
    division_type, _ = DivisionType.objects.get_or_create(
        code="department", defaults={"name": "department"}
    )
    return Division.objects.create(
        organization=org, type_code=division_type, name="Отдел Ф", code="EXPF-A"
    )


def make_employee(division, full_name="Сотрудник", rank_code=""):
    n = next(_iin_seq)
    return Employee.objects.create(
        iin=f"{n:012d}",
        full_name=full_name,
        rank_code=rank_code,
        position_code="",
        division=division,
        employment_status="WORKING",
    )


def make_status(employee, code, date_start, date_end, source="USER"):
    return EmployeeStatus.objects.create(
        employee_id=employee.id,
        status_type_code=code,
        date_start=date_start,
        date_end=date_end,
        source=source,
    )


def _expected_matrix(division, snapshot, staff=10, name="Отдел Ф"):
    """Независимое ожидание: derive_report от тех же входов, что и билдер,
    раскатанный в матрицу [Штат, Список, Вакансии, 12 колонок] по строкам
    (данные + ИТОГО)."""
    parsed = [
        {
            "employee_id": row["employee_id"],
            "status_type_code": row["status_type_code"],
            "date_start": date.fromisoformat(row["date_start"]),
            "date_end": date.fromisoformat(row["date_end"]),
        }
        for row in snapshot["rows"]
    ]
    roster_ids = [member["employee_id"] for member in snapshot["roster"]]
    report = derive_report(
        {division.id: roster_ids},
        parsed,
        {division.id: staff},
        D,
        division_names={division.id: name},
    )
    matrix = []
    for row in report.rows:
        numbers = [row.staff_total, row.list_total, row.vacancies]
        for key in DOCX_COLUMNS:
            numbers.append(row.attached if key == "ATTACHED" else row.columns[key])
        matrix.append(numbers)
    totals = report.totals
    totals_numbers = [totals.staff_total, totals.list_total, totals.vacancies]
    for key in DOCX_COLUMNS:
        totals_numbers.append(
            totals.attached if key == "ATTACHED" else totals.columns[key]
        )
    matrix.append(totals_numbers)
    return matrix


def _int(text):
    """Нормализация ATTACHED: ``+N`` → N."""
    return int(str(text).lstrip("+"))


def _matrix_docx(raw):
    table = Document(BytesIO(raw)).tables[0]
    matrix = []
    for row in table.rows[1:]:
        numbers = [_int(row.cells[column].text) for column in (2, 3, 4)]
        for offset in range(len(DOCX_COLUMNS)):
            cell = row.cells[_FIXED_HEAD_LEN + offset]
            numbers.append(_int(cell.paragraphs[0].text))
        matrix.append(numbers)
    return matrix


def _matrix_xlsx(raw):
    sheet = load_workbook(BytesIO(raw)).active
    matrix = []
    for sheet_row in range(3, sheet.max_row + 1):
        numbers = [
            _int(sheet.cell(row=sheet_row, column=column).value) for column in (3, 4, 5)
        ]
        for offset in range(len(DOCX_COLUMNS)):
            value = sheet.cell(row=sheet_row, column=6 + offset).value
            if isinstance(value, str):
                value = value.splitlines()[0]
            numbers.append(_int(value))
        matrix.append(numbers)
    return matrix


def _matrix_csv(raw):
    rows = list(csv.reader(StringIO(raw.decode("utf-8-sig")), delimiter=";"))
    return [[_int(field) for field in record[2:]] for record in rows[2:]]


def test_e2e_snapshot_to_four_formats_numbers_equal_derive(division):
    Rank.objects.create(code="CAPT", name="капитан")
    make_status(
        make_employee(division, "Отпускник Аскар", "CAPT"),
        "VACATION",
        date(2026, 7, 1),
        date(2026, 7, 15),
    )
    # Multi-код колонка: STUDY + CONFERENCE → ОДНА ячейка TRAINING (count 2).
    make_status(
        make_employee(division, "Студент Болат"),
        "STUDY",
        date(2026, 7, 5),
        date(2026, 7, 20),
    )
    make_status(
        make_employee(division, "Докладчик Виктор"),
        "CONFERENCE",
        date(2026, 7, 7),
        date(2026, 7, 9),
    )
    make_status(
        make_employee(division, "Дежурный Галым"),
        "DUTY",
        date(2026, 7, 8),
        date(2026, 7, 9),
    )
    make_status(
        make_employee(division, "Гость Дармен"),
        "ATTACHED",
        date(2026, 7, 1),
        date(2026, 8, 1),
    )
    make_employee(division, "Строевой Ержан")  # без факта → В строю
    make_employee(division, "Строевой Жанат")  # без факта → В строю

    snapshot = build_division_snapshot(division.id, D)
    data = build_expense_document(
        snapshot,
        D,
        staff_map={division.id: 10},
        division_names={division.id: "Отдел Ф"},
        division_id=division.id,
    )
    expected = _expected_matrix(division, snapshot)

    # Документные и числовая формы: поле-в-поле против derive, позиционно.
    assert _matrix_docx(generate_expense_docx(data)) == expected
    assert _matrix_xlsx(generate_expense_xlsx(data)) == expected
    assert _matrix_csv(generate_expense_csv(data)) == expected

    # PDF: каждое ожидаемое число присутствует токеном; ATTACHED — «+N».
    reader = PdfReader(BytesIO(generate_expense_pdf(data)))
    text = "\n".join(page.extract_text() for page in reader.pages)
    tokens = set(re.split(r"\D+", text))
    expected_numbers = {number for row in expected for number in row}
    missing = [number for number in expected_numbers if str(number) not in tokens]
    assert missing == []
    normalized = "".join(text.split())
    assert f"+{expected[0][3 + DOCX_COLUMNS.index('ATTACHED')]}" in normalized
    assert "ИТОГО" in normalized
    assert "ОтделФ" in normalized  # титул с именем подразделения
