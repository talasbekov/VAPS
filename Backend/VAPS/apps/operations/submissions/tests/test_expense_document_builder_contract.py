"""QA-судьи билдера данных .docx расхода (Story 6.3, bmad-qa-generate-e2e-tests).

Дополняют dev-сьют по мутационному принципу («какая правка кода не краснит
НИ ОДИН тест»): громкое падение Д11 на лишнем ключе ``staff_map``
(``(row,) = report.rows`` → ``rows[0]`` не краснил ничего), период члена —
из факта winner-КОДА (а не соседа по колонке), член ровно в ОДНОЙ ячейке при
фактах РАЗНЫХ колонок, и первый e2e хвостовых колонок Д3 (OTHER/BEFORE_DUTY/
PENDING + слияние GEV в «На дежурстве») с ВИДИМОЙ сходимостью Σ колонок ==
«По списку» числами из распарсенного документа (raison d'être Ловушки №4).

Конвенции dev-файла: ``django_db`` без ``transaction=True`` (teardown-ловушка
6.2), снапшот — ТОЛЬКО реальный ``build_division_snapshot`` (урок 5.4b);
dev-модуль ``test_expense_document.py`` не тронут. Пересечения фактов в
фикстурах легальны: HARD-эксклюзия (``excl_hard_status_overlap``) требует
двух HARD-кодов, здесь всегда максимум один.
"""

import itertools
from datetime import date
from io import BytesIO

import pytest
from docx import Document

from apps.core.models import Division, DivisionType, Employee, Organization
from apps.documents.generators import DOCX_COLUMNS, generate_expense_docx
from apps.operations.statuses.models import EmployeeStatus
from apps.operations.submissions.expense_document import build_expense_document
from apps.operations.submissions.services.snapshot import build_division_snapshot

pytestmark = pytest.mark.django_db

D = date(2026, 7, 8)
_iin_seq = itertools.count(800)

# Смещение статусной колонки: № | Управление | По штату | По списку | Вакансии.
_FIXED_HEAD_LEN = 5


def _col(key):
    return _FIXED_HEAD_LEN + DOCX_COLUMNS.index(key)


@pytest.fixture
def org():
    return Organization.objects.create(name="Орг", code="ORG-EXPQ")


@pytest.fixture
def division(org):
    division_type, _ = DivisionType.objects.get_or_create(
        code="department", defaults={"name": "department"}
    )
    return Division.objects.create(
        organization=org, type_code=division_type, name="Отдел К", code="EXPQ-A"
    )


def make_employee(division, full_name):
    n = next(_iin_seq)
    return Employee.objects.create(
        iin=f"{n:012d}",
        full_name=full_name,
        rank_code="",
        position_code="",
        division=division,
        employment_status="WORKING",
    )


def make_status(employee, code, date_start, date_end):
    return EmployeeStatus.objects.create(
        employee_id=employee.id,
        status_type_code=code,
        date_start=date_start,
        date_end=date_end,
        source="USER",
    )


def _build(division, snapshot, staff=10):
    return build_expense_document(
        snapshot,
        D,
        staff_map={division.id: staff},
        division_names={division.id: division.name},
        division_id=division.id,
    )


def test_extra_staff_map_division_fails_loud(division, org):
    # Д11: derive строит строку на КАЖДЫЙ ключ (employees | staff_map) —
    # лишний division в staff_map обязан падать громко на распаковке
    # ``(row,) = report.rows``, а не молча рендерить первую строку
    # (мутация ``rows[0]`` не краснила ни один тест).
    make_employee(division, "Строевой Ержан")
    other = Division.objects.create(
        organization=org,
        type_code=division.type_code,
        name="Чужой отдел",
        code="EXPQ-B",
    )
    snapshot = build_division_snapshot(division.id, D)
    with pytest.raises(ValueError):
        build_expense_document(
            snapshot,
            D,
            staff_map={division.id: 10, other.id: 5},
            division_names={division.id: division.name, other.id: other.name},
            division_id=division.id,
        )


def test_member_period_from_winner_code_fact_not_column_mate(division):
    # Один сотрудник, два факта ОДНОЙ колонки TRAINING разных кодов:
    # CONFERENCE (начат раньше) и STUDY (winner: приоритет 32 < 36). Период
    # члена — из факта winner-КОДА (07.07–10.07); фильтр «по колонке» утащил
    # бы min date_start соседа-CONFERENCE (01.07). Оба кода вне HARD —
    # пересечение легально; dev-судья min-date гоняет факты ОДНОГО кода.
    employee = make_employee(division, "Учащийся Болат")
    make_status(employee, "CONFERENCE", date(2026, 7, 1), date(2026, 7, 20))
    make_status(employee, "STUDY", date(2026, 7, 7), date(2026, 7, 10))

    data = _build(division, build_division_snapshot(division.id, D))
    cell = data.rows[0].cells["TRAINING"]
    assert cell.count == 1
    (member,) = cell.members
    assert member.full_name == "Учащийся Болат"
    assert member.date_start == date(2026, 7, 7)
    assert member.date_end == date(2026, 7, 10)


def test_facts_in_two_columns_member_only_in_winner_cell(division):
    # Факты в РАЗНЫХ колонках (SICK_LEAVE 10 побеждает DUTY 70): член ровно
    # в одной ячейке. Сбор «в каждую колонку действующего факта» разъехался
    # бы с derive (ON_DUTY 0 против 1 члена → cross-assert), но в dev-сьюте
    # у каждого сотрудника факты одной колонки — мутация была невидима.
    employee = make_employee(division, "Больной Санжар")
    make_status(employee, "SICK_LEAVE", date(2026, 7, 5), date(2026, 7, 15))
    make_status(employee, "DUTY", date(2026, 7, 8), date(2026, 7, 9))

    data = _build(division, build_division_snapshot(division.id, D))
    row = data.rows[0]
    assert row.cells["SICK"].count == 1
    (member,) = row.cells["SICK"].members
    assert member.full_name == "Больной Санжар"
    assert member.date_start == date(2026, 7, 5)
    assert member.date_end == date(2026, 7, 15)
    assert row.cells["ON_DUTY"].count == 0
    assert row.cells["ON_DUTY"].members == ()

    body = Document(BytesIO(generate_expense_docx(data))).tables[0].rows[1]
    assert body.cells[_col("SICK")].paragraphs[0].text == "1"
    assert body.cells[_col("ON_DUTY")].text == "0"


def test_tail_columns_visible_and_visible_sum_equals_list_total(division):
    # Первый e2e хвоста Д3: OTHER_ABSENCE→«Иное», BEFORE_DUTY→«Перед
    # дежурством», PENDING_CLARIFICATION→«Уточняется» доходят до ДОКУМЕНТА
    # с членами, GEV сливается с DUTY в одну ячейку «На дежурстве», и Σ
    # видимых статусных колонок == «По списку» числами ИЗ документа —
    # формула, ради которой хвост вообще виден (Ловушка №4).
    make_status(
        make_employee(division, "Иной Мади"),
        "OTHER_ABSENCE",
        date(2026, 7, 1),
        date(2026, 7, 15),
    )
    make_status(
        make_employee(division, "Готовящийся Нурлан"),
        "BEFORE_DUTY",
        date(2026, 7, 7),
        date(2026, 7, 9),
    )
    make_status(
        make_employee(division, "Неясный Олжас"),
        "PENDING_CLARIFICATION",
        date(2026, 7, 5),
        date(2026, 7, 10),
    )
    make_status(
        make_employee(division, "Дежурный Галым"),
        "DUTY",
        date(2026, 7, 8),
        date(2026, 7, 9),
    )
    make_status(
        make_employee(division, "Патрульный Ермек"),
        "GEV",
        date(2026, 7, 8),
        date(2026, 7, 9),
    )
    make_employee(division, "Строевой Жанат")  # без факта → В строю

    data = _build(division, build_division_snapshot(division.id, D))
    row = data.rows[0]
    assert row.cells["OTHER"].count == 1
    assert row.cells["BEFORE_DUTY"].count == 1
    assert row.cells["PENDING"].count == 1
    assert row.cells["ON_DUTY"].count == 2  # DUTY + GEV — одна колонка
    assert {m.full_name for m in row.cells["ON_DUTY"].members} == {
        "Дежурный Галым",
        "Патрульный Ермек",
    }

    document = Document(BytesIO(generate_expense_docx(data)))
    body = document.tables[0].rows[1]
    assert body.cells[_col("OTHER")].paragraphs[0].text == "1"
    assert "Иной Мади — 01.07.2026–15.07.2026" in [
        p.text for p in body.cells[_col("OTHER")].paragraphs
    ]
    assert body.cells[_col("BEFORE_DUTY")].paragraphs[0].text == "1"
    assert body.cells[_col("PENDING")].paragraphs[0].text == "1"
    assert body.cells[_col("ON_DUTY")].paragraphs[0].text == "2"
    visible = sum(
        int(body.cells[_col(key)].paragraphs[0].text)
        for key in DOCX_COLUMNS
        if key != "ATTACHED"
    )
    assert visible == int(body.cells[3].text) == 6
