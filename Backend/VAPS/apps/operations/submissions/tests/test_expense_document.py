"""Тесты билдера данных .docx расхода (Story 6.3, AC-3/4) + sync-гвард Д3.

Данные — прямые ``objects.create`` (зеркало test_snapshot_builder), снапшот —
ТОЛЬКО реальный ``build_division_snapshot`` (урок ревью 5.4b, руками не лепим;
единственное исключение — пустой снапшот-литерал как НЕГАТИВНЫЙ вход).
``django_db`` обычный, БЕЗ ``transaction=True`` (teardown-ловушка 6.2).

Sync-гвард дублированных ключей живёт здесь: submissions/tests легально видит
обе стороны границы (documents и operations); импорт operations из
documents/tests нарушал бы дух AST-гварда изоляции.
"""

import itertools
from datetime import date
from io import BytesIO

import pytest
from docx import Document

from apps.core.models import (
    Division,
    DivisionType,
    Employee,
    Organization,
    Rank,
)
from apps.documents.generators import DOCX_COLUMNS, generate_expense_docx
from apps.operations.statuses.models import EmployeeStatus
from apps.operations.statuses.services.strength_report import (
    ATTACHED_CODE,
    REPORT_COLUMNS,
    DivisionReportRow,
    StrengthReportResult,
    derive_report,
)
from apps.operations.submissions import expense_document
from apps.operations.submissions.expense_document import build_expense_document
from apps.operations.submissions.services.snapshot import build_division_snapshot

pytestmark = pytest.mark.django_db

D = date(2026, 7, 8)
_iin_seq = itertools.count(700)

# Смещение статусной колонки в таблице документа: № | Управление | По штату |
# По списку | Вакансии | 12 колонок DOCX_COLUMNS.
_FIXED_HEAD_LEN = 5


def _col(key):
    return _FIXED_HEAD_LEN + DOCX_COLUMNS.index(key)


@pytest.fixture
def org():
    return Organization.objects.create(name="Орг", code="ORG-EXPD")


@pytest.fixture
def division(org):
    division_type, _ = DivisionType.objects.get_or_create(
        code="department", defaults={"name": "department"}
    )
    return Division.objects.create(
        organization=org, type_code=division_type, name="Отдел Э", code="EXPD-A"
    )


def make_employee(division, full_name="Сотрудник", rank_code=""):
    n = next(_iin_seq)
    return Employee.objects.create(
        iin=f"{n:012d}",
        full_name=full_name,
        rank_code=rank_code,
        position_code="",
        division=division,
        employment_status="WORKING",
    )


def make_status(employee, code, date_start, date_end, source="USER"):
    return EmployeeStatus.objects.create(
        employee_id=employee.id,
        status_type_code=code,
        date_start=date_start,
        date_end=date_end,
        source=source,
    )


def _build(division, snapshot, staff=10, name="Отдел Э"):
    return build_expense_document(
        snapshot,
        D,
        staff_map={division.id: staff},
        division_names={division.id: name},
        division_id=division.id,
    )


def _derive_from_snapshot(division, snapshot, staff=10, name="Отдел Э"):
    """Независимое ожидание: тот же derive_report от тех же входов (AC-3)."""
    parsed = [
        {
            "employee_id": row["employee_id"],
            "status_type_code": row["status_type_code"],
            "date_start": date.fromisoformat(row["date_start"]),
            "date_end": date.fromisoformat(row["date_end"]),
        }
        for row in snapshot["rows"]
    ]
    roster_ids = [member["employee_id"] for member in snapshot["roster"]]
    return derive_report(
        {division.id: roster_ids},
        parsed,
        {division.id: staff},
        D,
        division_names={division.id: name},
    )


def test_docx_columns_cover_exactly_report_columns_plus_attached():
    # Sync-гвард Д3: осознанный дубль ключей через границу documents/operations
    # обязан покрывать РОВНО REPORT_COLUMNS + ATTACHED — дрейф ловится здесь.
    assert set(DOCX_COLUMNS) == set(REPORT_COLUMNS) | {ATTACHED_CODE}
    assert len(DOCX_COLUMNS) == len(REPORT_COLUMNS) + 1


def test_e2e_snapshot_to_docx_every_number_equals_derive(division):
    Rank.objects.create(code="CAPT", name="капитан")
    make_status(
        make_employee(division, "Отпускник Аскар", "CAPT"),
        "VACATION",
        date(2026, 7, 1),
        date(2026, 7, 15),
    )
    # Multi-code колонка (обязательный кейс): STUDY + CONFERENCE → ОДНА ячейка
    # TRAINING с count 2 и двумя членами.
    make_status(
        make_employee(division, "Студент Болат"),
        "STUDY",
        date(2026, 7, 5),
        date(2026, 7, 20),
    )
    make_status(
        make_employee(division, "Докладчик Виктор"),
        "CONFERENCE",
        date(2026, 7, 7),
        date(2026, 7, 9),
    )
    make_status(
        make_employee(division, "Дежурный Галым"),
        "DUTY",
        date(2026, 7, 8),
        date(2026, 7, 9),
    )
    make_status(
        make_employee(division, "Гость Дармен"),
        "ATTACHED",
        date(2026, 7, 1),
        date(2026, 8, 1),
    )
    make_employee(division, "Строевой Ержан")  # без факта → В строю
    make_employee(division, "Строевой Жанат")  # без факта → В строю

    snapshot = build_division_snapshot(division.id, D)
    data = _build(division, snapshot)
    expected = _derive_from_snapshot(division, snapshot)
    (expected_row,) = expected.rows

    # Билдер: каждое число == соответствующему полю derive_report — поле-в-поле.
    (row,) = data.rows
    assert row.name == expected_row.name
    assert row.staff_total == expected_row.staff_total
    assert row.list_total == expected_row.list_total
    assert row.vacancies == expected_row.vacancies
    for key in REPORT_COLUMNS:
        assert row.cells[key].count == expected_row.columns[key], key
    assert row.attached.count == expected_row.attached
    assert data.totals.staff_total == expected.totals.staff_total
    assert data.totals.list_total == expected.totals.list_total
    assert data.totals.vacancies == expected.totals.vacancies
    assert data.totals.columns == expected.totals.columns
    assert data.totals.attached == expected.totals.attached

    # Смысловая сверка состава (не только поле-в-поле): 2 в TRAINING, +1
    # прикомандированный вне «По списку», 2 без факта → В строю (расчёт derive:
    # 7 в ростере − 1 ATTACHED = 6 по списку).
    assert row.cells["TRAINING"].count == 2
    assert {m.full_name for m in row.cells["TRAINING"].members} == {
        "Студент Болат",
        "Докладчик Виктор",
    }
    assert row.attached.count == 1
    assert row.list_total == 6
    assert row.cells["IN_SERVICE"].count == 2

    # Документ: распарсить .docx и сверить все числа с derive.
    document = Document(BytesIO(generate_expense_docx(data)))
    table = document.tables[0]
    body = table.rows[1]
    assert body.cells[2].text == str(expected_row.staff_total)
    assert body.cells[3].text == str(expected_row.list_total)
    assert body.cells[4].text == str(expected_row.vacancies)
    for key in REPORT_COLUMNS:
        assert body.cells[_col(key)].paragraphs[0].text == str(
            expected_row.columns[key]
        ), key
    assert body.cells[_col("ATTACHED")].paragraphs[0].text == (
        f"+{expected_row.attached}"
    )
    totals_row = table.rows[-1]
    assert totals_row.cells[2].text == str(expected.totals.staff_total)
    assert totals_row.cells[3].text == str(expected.totals.list_total)
    assert totals_row.cells[4].text == str(expected.totals.vacancies)
    for key in REPORT_COLUMNS:
        assert totals_row.cells[_col(key)].text == str(expected.totals.columns[key])
    assert totals_row.cells[_col("ATTACHED")].text == f"+{expected.totals.attached}"
    # Ячейка TRAINING в документе: count + 2 строки членов с периодами.
    training_texts = [p.text for p in body.cells[_col("TRAINING")].paragraphs]
    assert training_texts[0] == "2"
    assert any("Студент Болат — 05.07.2026–20.07.2026" == t for t in training_texts)
    assert any("Докладчик Виктор — 07.07.2026–09.07.2026" == t for t in training_texts)
    # Заголовок несёт имя подразделения и дату.
    assert document.paragraphs[0].text == (
        "Отдел Э ЖЕКЕ ҚҰРАМЫНЫҢ САПТЫҚ ТІЗІМІ 08.07.2026 ЖЫЛҒЫ"
    )


def test_member_rank_resolved_and_period_from_winner_fact(division):
    Rank.objects.create(code="MAJ", name="майор")
    employee = make_employee(division, "Отпускник Аскар", "MAJ")
    make_status(employee, "VACATION", date(2026, 7, 2), date(2026, 7, 12))

    data = _build(division, build_division_snapshot(division.id, D))
    (member,) = data.rows[0].cells["VACATION"].members
    assert member.rank == "майор"
    assert member.full_name == "Отпускник Аскар"
    assert member.date_start == date(2026, 7, 2)
    assert member.date_end == date(2026, 7, 12)


def test_member_period_multiple_winner_facts_takes_min_date_start(division):
    # DUTY — вне HARD_STATUS_TYPE_CODES, два пересекающихся факта легальны.
    employee = make_employee(division, "Дежурный Галым")
    make_status(employee, "DUTY", date(2026, 7, 5), date(2026, 7, 12))
    make_status(employee, "DUTY", date(2026, 7, 1), date(2026, 7, 10))

    data = _build(division, build_division_snapshot(division.id, D))
    cell = data.rows[0].cells["ON_DUTY"]
    assert cell.count == 1
    (member,) = cell.members
    assert member.date_start == date(2026, 7, 1)  # Д7: min по date_start
    assert member.date_end == date(2026, 7, 10)


def test_member_period_equal_start_tie_breaks_by_min_date_end(division):
    # Ревью 6.3 (Edge-слой): два факта winner-кода с ОДНИМ date_start — тай
    # ломается по min date_end детерминированно, а не порядком строк снапшота
    # (длинный факт создан ПЕРВЫМ: выбор «первый в row-order» дал бы 15.07).
    employee = make_employee(division, "Дежурный Галым")
    make_status(employee, "DUTY", date(2026, 7, 1), date(2026, 7, 15))
    make_status(employee, "DUTY", date(2026, 7, 1), date(2026, 7, 10))

    data = _build(division, build_division_snapshot(division.id, D))
    cell = data.rows[0].cells["ON_DUTY"]
    assert cell.count == 1
    (member,) = cell.members
    assert member.date_start == date(2026, 7, 1)
    assert member.date_end == date(2026, 7, 10)


def test_half_open_interval_fact_ending_on_business_date_not_acting(division):
    # Полуоткрытый [start, end): факт с date_end == business_date уже НЕ
    # действует (класс ошибок ARCH-DATA-023, Ловушка №8).
    employee = make_employee(division, "Вернувшийся Ержан")
    make_status(employee, "VACATION", date(2026, 7, 1), D)

    data = _build(division, build_division_snapshot(division.id, D))
    row = data.rows[0]
    assert row.cells["VACATION"].count == 0
    assert row.cells["VACATION"].members == ()
    assert row.cells["IN_SERVICE"].count == 1


def test_roster_member_without_fact_counts_in_service_without_members(division):
    make_employee(division, "Строевой Жанат")

    data = _build(division, build_division_snapshot(division.id, D))
    cell = data.rows[0].cells["IN_SERVICE"]
    assert cell.count == 1
    assert cell.members == ()


def test_event_assignment_counts_in_service_without_members(division):
    # EVENT_ASSIGNMENT маппится в колонку IN_SERVICE (strength_report.py):
    # счётчик растёт, члены НЕ собираются даже при живом факте (AC-4).
    employee = make_employee(division, "Мероприятие Канат")
    make_status(employee, "EVENT_ASSIGNMENT", date(2026, 7, 1), date(2026, 7, 20))

    data = _build(division, build_division_snapshot(division.id, D))
    cell = data.rows[0].cells["IN_SERVICE"]
    assert cell.count == 1
    assert cell.members == ()


def test_attached_member_collected_and_out_of_list_total(division):
    employee = make_employee(division, "Гость Дармен")
    make_status(employee, "ATTACHED", date(2026, 7, 1), date(2026, 8, 1))

    data = _build(division, build_division_snapshot(division.id, D))
    row = data.rows[0]
    assert row.attached.count == 1
    assert [m.full_name for m in row.attached.members] == ["Гость Дармен"]
    assert row.list_total == 0  # прикомандированный вне «По списку»


def test_cross_assert_count_vs_members_via_monkeypatched_derive(division, monkeypatch):
    # Честным входом счётчики и члены не развести (один снапшот) —
    # санкционированная техника: подделать derive_report в модуле билдера.
    employee = make_employee(division, "Дежурный Галым")
    make_status(employee, "DUTY", date(2026, 7, 1), date(2026, 7, 10))
    snapshot = build_division_snapshot(division.id, D)

    def tampered_derive(*args, **kwargs):
        real = derive_report(*args, **kwargs)
        (row,) = real.rows
        columns = dict(row.columns)
        columns["ON_DUTY"] += 1  # count разъехался со списком членов
        return StrengthReportResult(
            business_date=real.business_date,
            rows=[
                DivisionReportRow(
                    division_id=row.division_id,
                    name=row.name,
                    staff_total=row.staff_total,
                    list_total=row.list_total,
                    vacancies=row.vacancies,
                    columns=columns,
                    attached=row.attached,
                )
            ],
            totals=real.totals,
            violations=real.violations,
            warnings=real.warnings,
        )

    monkeypatch.setattr(expense_document, "derive_report", tampered_derive)
    with pytest.raises(AssertionError):
        _build(division, snapshot)


def test_empty_snapshot_literal_builds_zero_row_without_crash(division):
    # Литерал легален: негативный вход, не позитивная фикстура (урок 5.4b).
    data = _build(division, {"roster": [], "rows": []})
    (row,) = data.rows
    assert row.list_total == 0
    assert row.staff_total == 10
    assert row.vacancies == 10
    assert all(row.cells[key].count == 0 for key in REPORT_COLUMNS)
    assert all(row.cells[key].members == () for key in REPORT_COLUMNS)
    assert row.attached.count == 0
    generate_expense_docx(data)  # и рендер не падает


def test_unknown_status_type_code_raises_value_error(division):
    # STOP-семантика (AC-3): исключения derive НЕ глотаются.
    employee = make_employee(division, "Загадка Мурат")
    make_status(employee, "MYSTERY", date(2026, 7, 1), date(2026, 7, 10))

    with pytest.raises(ValueError):
        _build(division, build_division_snapshot(division.id, D))
