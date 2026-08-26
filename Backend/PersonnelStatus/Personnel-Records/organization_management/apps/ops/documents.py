"""Документы ОМ: шаблон `.docx` → подстановка → PDF (Plane №157, шаг ПД-1).

ЗАЧЕМ ИМЕННО ТАК. Заказчик сказал дважды: «Документы должны выглядеть В
ТОЧНОСТИ как ворд формате которые я дал тебе» и «выгрузка должна быть в пдф, а
выглядеть как ворд файл». Значит документ не рисуется заново, а БЕРЁТСЯ
готовым: образец заказчика — это шаблон, в нём меняются только значения.

Почему не сборка с нуля (`reportlab`). Она даёт похожую структуру и непохожий
вид. В образце «Сводные данные» есть флаг страны картинкой, две фотографии,
объединённые ячейки, заливка заголовков, подчёркивания внутри ячеек и курсив в
скобках — перерисовать это «в точности» нельзя, а заказчик просил именно
в точности. Генератор таблиц остаётся для документов, у которых образца НЕТ.

Почему это дёшево. `LibreOffice` и `python-docx` уже стоят: ставить нечего.
Замер конвертации образца — 0,48 с (два прогона подряд одинаково).

ЧЕГО ЗДЕСЬ НЕТ И ПОЧЕМУ. Шаблонов конкретных документов: они приезжают своими
шагами (ПД-2 и дальше). Здесь только конвейер и его правила.
"""
import os
import re
import shutil
import subprocess
import tempfile

from organization_management.apps.operations.exceptions import DomainError

#: Место подстановки в шаблоне: `{{ключ}}`. Фигурные скобки, а не `$ключ` и не
#: `%ключ%`: в текстах документов встречаются и проценты, и доллары, а двойная
#: фигурная скобка в делопроизводстве не встречается вовсе.
PLACEHOLDER = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")

#: Потолок ожидания конвертации. LibreOffice на образце укладывается в 0,5 с;
#: 60 с — это «что-то пошло не так», а не рабочий случай.
CONVERT_TIMEOUT_SECONDS = 60


def _fill_paragraph(paragraph, values):
    """Подставить значения в один параграф.

    ТОНКОСТЬ, БЕЗ КОТОРОЙ ПОДСТАНОВКА МОЛЧА НЕ РАБОТАЕТ. Word режет текст на
    «прогоны» (runs) по границам форматирования, и делает это непредсказуемо:
    `{{country}}` легко оказывается разложенным на `{{cou`, `ntry`, `}}` — по
    прогону на кусок. Замена по каждому прогону отдельно не найдёт ничего и
    не сообщит об этом: документ выйдет с местами подстановки вместо значений.

    Поэтому текст параграфа собирается целиком, заменяется, и результат
    кладётся в ПЕРВЫЙ прогон, а остальные очищаются. Форматирование первого
    прогона при этом сохраняется — а внутри одного места подстановки его и не
    бывает разным.
    """
    full = "".join(run.text for run in paragraph.runs)
    if "{{" not in full:
        return
    filled = PLACEHOLDER.sub(
        lambda m: str(values.get(m.group(1), m.group(0))), full
    )
    if filled == full:
        return
    paragraph.runs[0].text = filled
    for run in paragraph.runs[1:]:
        run.text = ""


def _fill_document(document, values):
    """Подстановка по ВСЕМУ документу: тело, таблицы, колонтитулы.

    Колонтитулы перечислены явно: в образцах заказчика дата и время среза
    стоят в шапке («проект на 22.04.2026 г. время 08:00»), а обход `document`
    их не видит — это отдельная часть файла.
    """
    for paragraph in document.paragraphs:
        _fill_paragraph(paragraph, values)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _fill_paragraph(paragraph, values)
    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                _fill_paragraph(paragraph, values)


def unresolved_placeholders(document):
    """Места подстановки, оставшиеся в документе после заполнения.

    Нужна пробам и самой выгрузке: документ с `{{employee}}` вместо фамилии
    выглядит как готовый и уходит наружу как готовый. Молчать об этом нельзя.
    """
    found = set()
    for paragraph in document.paragraphs:
        found.update(PLACEHOLDER.findall(paragraph.text))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                found.update(PLACEHOLDER.findall(cell.text))
    return sorted(found)


def fill_template(template_path, values):
    """Заполнить шаблон и вернуть путь к заполненному `.docx` во временной
    папке. Возвращается ПАРА (путь, оставшиеся места подстановки) — решение,
    что делать с недозаполненным документом, принимает вызывающий."""
    from docx import Document

    if not os.path.exists(template_path):
        raise DomainError(
            "DOCUMENT_TEMPLATE_MISSING", 500,
            detail={"template": [template_path]},
            message="Шаблон документа не найден.",
        )
    try:
        document = Document(template_path)
    except Exception as error:
        # Битый или обрезанный файл. Пойман на образце заказчика
        # «01 Сводные данные РЭС 22.04.docx»: у него НЕТ конца zip-архива
        # (`EOCD`) и центрального каталога — файл недокачан. `file` при этом
        # честно говорит «Microsoft Word 2007+», потому что смотрит только на
        # первые байты, а LibreOffice отвечает невнятным «source file could
        # not be loaded».
        #
        # Отказ обязан называть ФАЙЛ и ПРИЧИНУ: иначе на демонстрации это
        # выглядит как «выгрузка не работает», и чинить будут конвейер вместо
        # того, чтобы перевыслать документ.
        raise DomainError(
            "DOCUMENT_TEMPLATE_BROKEN", 500,
            detail={"template": [f"{template_path}: {error}"]},
            message="Шаблон документа повреждён и не читается.",
        )
    _fill_document(document, values)
    left = unresolved_placeholders(document)
    handle, filled_path = tempfile.mkstemp(suffix=".docx")
    os.close(handle)
    document.save(filled_path)
    return filled_path, left


def docx_to_pdf(docx_path):
    """Конвертировать `.docx` в PDF и вернуть байты.

    СВОЙ ПРОФИЛЬ НА КАЖДЫЙ ВЫЗОВ (`-env:UserInstallation`). LibreOffice держит
    в профиле блокировку, и две одновременные конвертации с общим профилем
    мешают друг другу: вторая либо ждёт, либо падает. Выгрузку могут нажать
    два человека разом — общий профиль сделал бы это гонкой.
    """
    if shutil.which("soffice") is None:
        # Системная зависимость, которой нет в `requirements.txt`. Молчаливый
        # откат «отдать .docx вместо PDF» запрещён: заказчик просил PDF, и
        # подмена формата — не помощь, а сюрприз.
        raise DomainError(
            "PDF_CONVERTER_MISSING", 500,
            detail={"converter": ["LibreOffice (soffice) не установлен."]},
            message="Конвертер PDF недоступен.",
        )
    workdir = tempfile.mkdtemp()
    try:
        result = subprocess.run(
            [
                "soffice",
                f"-env:UserInstallation=file://{workdir}/profile",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", workdir,
                docx_path,
            ],
            capture_output=True,
            timeout=CONVERT_TIMEOUT_SECONDS,
        )
        produced = os.path.join(
            workdir,
            os.path.splitext(os.path.basename(docx_path))[0] + ".pdf",
        )
        if not os.path.exists(produced):
            raise DomainError(
                "PDF_CONVERSION_FAILED", 500,
                detail={"stderr": [result.stderr.decode("utf-8", "ignore")[:400]]},
                message="Документ не удалось собрать в PDF.",
            )
        with open(produced, "rb") as handle:
            return handle.read()
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


#: Форматы выгрузки. DOCX — то, что просил заказчик: образцы это РАБОЧИЕ
#: бланки Word, их дозаполняют руками после выгрузки. PDF рядом нужен, когда
#: документ идут печатать или отправлять и правок в нём не ждут.
FORMATS = ("docx", "pdf")

CONTENT_TYPES = {
    "docx": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ),
    "pdf": "application/pdf",
}


def emit(filled_path, fmt):
    """Заполненный `.docx` → байты в запрошенном формате.

    Одна точка на все пять сборщиков: до неё каждый звал `docx_to_pdf`
    напрямую, и формат был вшит в сборщик. Спрашивать формат у каждого
    сборщика по отдельности значило бы получить пять разных ответов на один
    вопрос — так уже вышло с их подписями (см. `documents_registry`).
    """
    if fmt not in FORMATS:
        raise DomainError(
            "VALIDATION_ERROR", 400,
            detail={"format": ["Формат бывает: " + ", ".join(FORMATS)]},
            message="Проверьте заполнение формы.",
        )
    if fmt == "docx":
        with open(filled_path, "rb") as handle:
            return handle.read()
    return docx_to_pdf(filled_path)


def render_docx_from_template(template_path, values, *, allow_unresolved=False):
    """Шаблон + значения → байты `.docx`.

    ЗАЧЕМ ОТДЕЛЬНО ОТ PDF. Заказчик просил документы «в таком же формате», а
    образцы — рабочие бланки WORD: их дозаполняют руками после выгрузки. PDF
    этого не даёт. Конвейер и так заполняет `.docx` и лишь потом зовёт
    LibreOffice — значит нужный формат уже собран, и не отдавать его было бы
    решением за заказчика.

    PDF рядом НЕ снимается: он нужен, когда документ идут печатать или
    отправлять, и правки в нём не ждут.
    """
    filled_path, _left = _fill_or_fail(template_path, values, allow_unresolved)
    try:
        return emit(filled_path, "docx")
    finally:
        _drop_temp(filled_path, template_path)


def _drop_temp(filled_path, template_path):
    """Убрать ВРЕМЕННУЮ копию шаблона.

    Сравнение с исходным путём — не паранойя. `fill_template` всегда отдаёт
    копию, но функции сборки удаляют то, что она вернула, в `finally`; стоит
    кому-то передать сюда сам шаблон, и удалится ОРИГИНАЛ — тихо, без ошибки,
    и заметится это только когда следующая сборка не найдёт бланк. Так и
    случилось при проверке мутацией (Plane №156): шаблон «Сводных данных»
    исчез с диска, и шесть проб покраснели вдалеке от причины.
    """
    if os.path.abspath(filled_path) == os.path.abspath(template_path):
        return
    try:
        os.unlink(filled_path)
    except OSError:
        pass


def _fill_or_fail(template_path, values, allow_unresolved):
    """Заполнить шаблон, отказав на недозаполненном. Общее у обоих форматов:
    правило «не выпускать документ с `{{...}}`» не должно зависеть от того,
    в чём его попросили."""
    filled_path, left = fill_template(template_path, values)
    if left and not allow_unresolved:
        _drop_temp(filled_path, template_path)
        raise DomainError(
            "DOCUMENT_INCOMPLETE", 500,
            detail={"placeholders": left},
            message="Документ заполнен не полностью.",
        )
    return filled_path, left


def render_pdf_from_template(template_path, values, *, allow_unresolved=False):
    """Шаблон + значения → байты PDF.

    `allow_unresolved=False` по умолчанию: недозаполненный документ НЕ уходит
    наружу. Пустое значение и незаполненное место — разные вещи: пустое поле
    заказчик прочтёт как «сведений нет», а `{{employee}}` — как поломку, и он
    будет прав.
    """
    filled_path, _ = _fill_or_fail(template_path, values, allow_unresolved)
    try:
        return docx_to_pdf(filled_path)
    finally:
        _drop_temp(filled_path, template_path)
