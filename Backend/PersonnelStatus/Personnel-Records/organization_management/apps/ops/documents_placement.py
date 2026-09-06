"""Документ расстановки: посты, секторы, назначенные (Plane №156, шаг «ПД-6»).

🔴 ДОПУЩЕНИЕ, О КОТОРОМ НАДО ЗНАТЬ ДО ЧТЕНИЯ КОДА. У остальных документов
цепочки шаблон снят С ОБРАЗЦА заказчика, и они выглядят «в точности как ворд».
С расстановкой так не вышло, и вот почему:

* `Расстановка Алем Ай 25 ОБРАЗЕЦ.doc` — ФАЙЛ БИТЫЙ: заголовок OLE на месте,
  но потока `WordDocument` внутри НЕТ, и ни LibreOffice, ни Word его не
  откроют (соседний `.doc` из той же папки конвертируется, то есть дело в
  файле, а не в формате);
* `Общая расстановка РЭС.DOCX` — рукодельная вёрстка ПОД КОНКРЕТНОЕ
  мероприятие: 14 таблиц разной формы (6 и 4 колонки), без строк-заголовков,
  на казахском, с грифом и подписями руководства. Это не бланк, который
  заполняют данными, а документ, который каждый раз собирают заново.

Поэтому таблица здесь собрана ПО ПОЛЯМ ШАГА («посты, секторы, назначенные»), а
ВЁРСТКА — рамки, заливка заголовка, шрифты — взята у бюллетеня, то есть у
образца заказчика. Как только придёт годный образец расстановки, шаблон
меняется на него, а код остаётся: он отдаёт строки, а не рисует.

Данные — из расчёта постов рекогносцировки и назначений расстановки самого
мероприятия: второй источник тех же сведений разошёлся бы с доской подбора.
"""
import datetime as dt
import os

from organization_management.apps.operations.clock import Clock
from organization_management.apps.operations.exceptions import DomainError
from organization_management.apps.operations.models_event import OpsSecurityEvent
from organization_management.apps.ops.document_tables import fill_table_rows
from organization_management.apps.ops.documents import emit, fill_template, stamp_draft

TEMPLATE = os.path.join(
    os.path.dirname(__file__), "document_templates", "placement.docx"
)


def _document_target(event, visit_object_id):
    """Объект посещения, чей документ собирается (Plane №411, Ш-5).

    Спецификация `[МД-04]`: документ «Расстановка сил» — СВОЙ У ОБЪЕКТА. Общий
    документ мероприятия с двумя объектами сваливал посты двух разных мест в
    одну таблицу, и подписывать его было нечем: согласуют объект.

    Объектов нет вовсе — собирается ВЕСЬ расчёт мероприятия, как и до этого
    шага. Это не поблажка, а сохранение живых данных: у таких ОМ расчёт лежит
    в мероприятии, и отказ вместо документа отнял бы у них то, что работало.
    """
    from organization_management.apps.ops import security_events

    if not event.visit_objects.exists():
        return None
    return security_events.pick_visit_object(
        event,
        visit_object_id,
        no_objects="",  # недостижимо: список проверен строкой выше
        ambiguous=(
            "У мероприятия несколько объектов посещения — выберите, чей "
            "документ «Расстановка сил» собрать: документ принадлежит объекту."
        ),
    )


def _assigned_names(event, post_id):
    """Кто стоит на посту — фамилиями и позывными, столбиком.

    Имена берутся из САМОГО назначения (`placement_assignments`), а не из
    кадров по идентификатору: в назначении лежит имя на момент расстановки, и
    оно должно остаться таким же в документе, даже если человека потом
    переименовали или уволили.
    """
    names = []
    for row in event.placement_assignments or []:
        if str(row.get("postId")) != str(post_id):
            continue
        name = str(row.get("employeeName") or "").strip()
        callsign = str(row.get("callsign") or "").strip()
        names.append(" ".join(part for part in (name, callsign) if part) or "—")
    return "\n".join(names)


def placement_rows(event, visit=None):
    """Строки документа: по одной на пост расчёта, в порядке секторов.

    Объект назван — только ЕГО посты (`visit_object_posts`, тот же разрез, что
    у согласования и у экрана этапа).
    """
    if visit is None:
        posts = event.recon_sector_posts or []
    else:
        from organization_management.apps.ops import security_events

        posts = security_events.visit_object_posts(event, visit)
    rows = []
    for post in posts:
        post_id = post.get("id")
        rows.append(
            {
                "sector": str(post.get("sector") or ""),
                "post": str(post.get("post") or post.get("name") or ""),
                "task": str(post.get("task") or ""),
                # Смена у поста появилась шагом Plane №123; у ОМ, заведённых
                # раньше, её нет — и пустая ячейка честнее выдуманной «дневной».
                "shift": str(post.get("shift") or ""),
                "need": post.get("need") if post.get("need") is not None else "",
                "assigned": _assigned_names(event, post_id),
            }
        )
    return rows


def signature_lines(visit):
    """Строки подвала по подписям маршрута объекта; без подписей — пусто."""
    if visit is None:
        return []
    lines = []
    for item in visit.approval_route or []:
        signature = item.get("signature") if item.get("status") == "APPROVED" else None
        if not signature:
            continue
        signed_at = signature.get("signedAt") or item.get("decidedAt") or ""
        try:
            moment = dt.datetime.fromisoformat(signed_at)
            when = moment.strftime("%d.%m.%Y %H:%M")
        except ValueError:
            when = signed_at
        position = signature.get("position") or item.get("position") or ""
        lines.append(
            "Согласовано: "
            f"{signature.get('fullName') or item.get('name', '')}, {position}, "
            f"{when}, версия {signature.get('versionNumber', visit.document_version)}"
        )
    return lines


def _is_draft(event, visit):
    """Документ ещё не согласован — на бумаге он проект (`[СОГ-03]`).

    🔴 У ОМ БЕЗ ОБЪЕКТОВ ПОСЕЩЕНИЯ СМОТРИМ НА САМО МЕРОПРИЯТИЕ (Plane №638).
    Раньше `visit is None` означало «проект» безусловно, и такой документ
    штамповался «ПРОЕКТ» НАВСЕГДА — даже согласованный. При этом экран
    предупреждение о черновике прятал: `ApprovalStage` читает `view.status`,
    который при пустом объекте падает на `event.approvalStatus`. Экран обещал
    чистый документ, сервер отдавал черновик — и спорить с бумагой человеку
    было нечем.

    Мероприятия без объектов посещения `_document_target` поддерживает
    НАМЕРЕННО (см. его докстринг: у таких ОМ расчёт лежит в самом мероприятии),
    поэтому «нет объекта» — не повод считать документ вечным проектом, а повод
    спросить о согласовании ту сущность, которая его несёт.
    """
    source = event if visit is None else visit
    return source.approval_status != "APPROVED"


def render_placement(event_code, as_of=None, fmt="pdf", visit_object_id=None):
    """Байты расстановки ОБЪЕКТА посещения; `fmt` — «docx» либо «pdf»."""
    from docx import Document

    event = OpsSecurityEvent.objects.filter(code=event_code).first()
    if event is None:
        raise DomainError(
            "ENTITY_NOT_FOUND",
            404,
            detail={"code": [str(event_code)]},
            message="Мероприятие не найдено.",
        )
    visit = _document_target(event, visit_object_id)
    moment = as_of or Clock.now()
    if isinstance(moment, dt.date) and not isinstance(moment, dt.datetime):
        moment = dt.datetime.combine(moment, dt.time(8, 0))
    # Объект и версия дописываются В СТРОКУ ЗАГОЛОВКА, а не своими метками:
    # шаблон достался от бюллетеня (см. докстринг модуля) и новых мест под
    # текст не имеет, а править чужую вёрстку ради двух слов — менять то, что
    # заказчик утверждал, ради того, чего он ещё не видел.
    #
    # «версия N» печатается ТОЛЬКО у объекта, чей документ уже уходил
    # согласующим: «версия 0» на бумаге читалась бы как номер, а означает
    # «не отправляли».
    title = f"{event.code} — {event.title}"
    if visit is not None:
        title += f" · объект: {visit.object_name}"
        if visit.document_version:
            title += f" · версия документа {visit.document_version}"
    values = {
        "event": title,
        "as_of": (
            f"{moment.strftime('%H:%M')} ч. "
            f"{moment.day:02d}.{moment.month:02d}.{moment.year} года"
        ),
    }
    filled_path, _left = fill_template(TEMPLATE, values)
    try:
        document = Document(filled_path)
        fill_table_rows(document.tables[0], placement_rows(event, visit))
        # Подвал подписей (`[СОГ-10]`, Plane №429): «Согласовано: ФИО,
        # должность, ДД.ММ.ГГГГ ЧЧ:ММ, версия N» — по строке на подпись.
        for line in signature_lines(visit):
            document.add_paragraph(line)
        # `[ОЗН-07]` (Plane №437): после завершения ознакомления документ
        # получает приложение «Лист ознакомления».
        #
        # 🔴 УСЛОВИЯ «есть объект посещения» ЗДЕСЬ БОЛЬШЕ НЕТ (Plane №697).
        # Мероприятия без объектов достижимы, и `_document_target` возвращает
        # для них `None` СОЗНАТЕЛЬНО — собирается весь расчёт мероприятия.
        # Приложение же под это условие не попадало, и `render_case` печатал
        # лист, а `render_placement` для того же ОМ — нет: два документа об
        # одном мероприятии расходились. Сборщик строк листа `visit=None`
        # умеет с самого начала — берёт все посты расчёта.
        #
        # Порог стадии остался на месте: приложение появляется ПОСЛЕ
        # завершения ознакомления, и без объектов это правило то же.
        from organization_management.apps.ops import documents_case

        # Здесь `visit` — ОБЪЕКТ ПОСЕЩЕНИЯ (Plane №907): у именованного
        # параметра то же имя, что у типа, и вызов теперь сам говорит, что
        # именно в него кладут.
        if documents_case.acknowledgement_completed(event, visit_object=visit):
            documents_case.append_acknowledgement_sheet(
                document, event, visit_object=visit
            )
        document.save(filled_path)
        payload = emit(filled_path, fmt)
        # `[СОГ-03]` (Plane №430): «Скачать PDF» доступна всегда, до
        # согласования — с водяным знаком «Проект». Проект — пока объект не
        # согласован (статус согласования ОБЪЕКТА, а не мероприятия: у
        # соседнего объекта своя подпись); у ОМ без объектов посещения спросить
        # не у кого, и там берётся статус самого мероприятия — Plane №638.
        # DOCX без знака: его дозаполняют руками, и знак в правимом файле — не
        # знак, а помеха.
        if fmt == "pdf" and _is_draft(event, visit):
            payload = stamp_draft(payload)
        return payload
    finally:
        try:
            os.unlink(filled_path)
        except OSError:
            pass
