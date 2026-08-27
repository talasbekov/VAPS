"""Бланк «Список броней в ГОН» из образца заказчика (Plane №216).

Вёрстка образца сохраняется — шапка, подписи колонок, ширины, — а ВСЕ строки
данных снимаются. Это не аккуратность, а требование: в образце
`04 Список броней в ГОН 05.91.2025 г. обнов.docx` стоят НАСТОЯЩИЕ
государственные номера машин, возящих охраняемых лиц. Ровно такая утечка уже
была найдена в этом репозитории (Plane №164, 73 позывных и 66 фамилий), и
повторять её нельзя.

Запускается руками, результат коммитится:
    .venv/bin/python organization_management/apps/ops/document_templates/build_vehicles_template.py
"""
import os
import re
import zipfile

from docx import Document

SRC = (
    "/home/erda/Музыка/Smart Josparlau/docs/PersonnelStatus/"
    "04 Список броней в ГОН 05.91.2025 г. обнов.docx"
)
DST = os.path.join(os.path.dirname(__file__), "vehicles_armored.docx")

#: Места подстановки строки таблицы — по колонкам образца.
ROW_PLACEHOLDERS = [
    "{{no}}",
    "{{brand}}",
    "{{body_class}}",
    "{{production_year}}",
    "{{plate}}",
    "{{armor_class}}",
    "{{deployment}}",
    "{{note}}",
]


def _set_cell(cell, text):
    """Текст ячейки, не теряя её оформления: пишем в первый прогон."""
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def main():
    document = Document(SRC)

    # Шапка: дата среза становится местом подстановки. Дату конкретного
    # документа в бланке оставлять нельзя — бланк не про один день.
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith("проект на"):
            _set_cell_paragraph(paragraph, "проект на {{as_of_date}} г.")

    table = document.tables[0]
    # Строка 0 — подписи колонок, остаётся как есть.
    #
    # Строка 1 образца — СЛУЖЕБНАЯ: «900» во всех ячейках, и все восемь ячеек
    # там — ОДНА склеенная. Строкой-образцом она не годится: заполнение
    # восьми колонок легло бы в одну (проверено — все восемь получили
    # последнее место подстановки). Она снимается вместе с данными.
    #
    # Образцом становится ПЕРВАЯ СТРОКА ДАННЫХ: у неё восемь настоящих ячеек
    # и то самое оформление, которым набран весь список.
    template_row = table.rows[2]
    if len(template_row.cells) != len(ROW_PLACEHOLDERS):
        raise SystemExit(
            f"в строке образца {len(template_row.cells)} ячеек, "
            f"а колонок {len(ROW_PLACEHOLDERS)} — образец изменился"
        )
    for cell, placeholder in zip(template_row.cells, ROW_PLACEHOLDERS):
        _set_cell(cell, placeholder)

    # Служебная строка «900» и все остальные строки данных снимаются.
    table.rows[1]._element.getparent().remove(table.rows[1]._element)
    for row in list(table.rows[2:]):
        row._element.getparent().remove(row._element)

    document.save(DST)
    _scrub_properties(DST)
    print(f"шаблон записан: {DST}")
    print(f"строк в таблице: {len(Document(DST).tables[0].rows)} (шапка + образец)")


#: Свойства файла, которые обязаны быть пусты у БЛАНКА. Бланк — форма, а не
#: документ: автора, организации и темы у него нет, и любое значение здесь —
#: след того, у кого форму сняли. Сторож `test_no_template_carries_filled_in_
#: properties` проверяет ровно эти поля.
PROPERTY_FIELDS = (
    "dc:creator",
    "cp:lastModifiedBy",
    "dc:title",
    "dc:subject",
    "dc:description",
    "cp:keywords",
    "Company",
    "Manager",
)


def _scrub_properties(path):
    """Вычистить свойства файла, унаследованные от образца.

    `python-docx` сохраняет `docProps/` образца как есть, включая автора и
    организацию заказчика. Пишем архив заново: правка на месте в zip
    невозможна.
    """
    source = zipfile.ZipFile(path)
    items = [(item, source.read(item.filename)) for item in source.infolist()]
    source.close()
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as out:
        for item, data in items:
            if item.filename.startswith("docProps/"):
                text = data.decode("utf-8", errors="ignore")
                for field in PROPERTY_FIELDS:
                    text = re.sub(
                        rf"<{field}([^>]*)>.*?</{field}>",
                        rf"<{field}\1></{field}>",
                        text,
                        flags=re.S,
                    )
                data = text.encode("utf-8")
            out.writestr(item, data)


def _set_cell_paragraph(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


if __name__ == "__main__":
    main()
