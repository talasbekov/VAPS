"""Бланк «Общая расстановка» из образца заказчика.

РЕШЕНИЕ ЗАКАЗЧИКА 27.08.2026: «Удали такие слова как Құпия и сделай выгрузку
точно такого же файла, обезлич все внутри».

Значит вёрстка образца сохраняется ЦЕЛИКОМ — 14 таблиц, колонки, шрифты,
казахские подписи, — а внутри не остаётся ни одной настоящей фамилии,
позывного, даты и грифа.
"""
import io, re, zipfile
from docx import Document
from PIL import Image, ImageDraw

SRC = "/home/erda/Музыка/Smart Josparlau/docs/PersonnelStatus/Общая расстановка РЭС.DOCX"
DST = "organization_management/apps/ops/document_templates/placement_full.docx"

# Пометки и блок утверждения — СНИМАЮТСЯ по прямому указанию заказчика.
# Гриф секретности в бланке, который живёт в репозитории, недопустим сам по
# себе; подпись руководителя — это утверждение конкретного документа, а не
# форма.
DROP_LINES = {
    "Құпия", "«БЕКІТЕМІН»", "БЕКІТЕМІН",
    "Қазақстан Республикасы", "Мемлекеттік күзет қызметі",
    "бастығының орынбасары", "генерал-майор",
}
#: Звания в строке подписи. Подпись — это утверждение КОНКРЕТНОГО документа,
#: а не часть формы: в бланке её быть не должно.
RANKS = ("полковник", "подполковник", "генерал-майор", "генерал-лейтенант",
         "майор", "капитан")

DROP_PATTERNS = [
    re.compile(r"^№\s*\d+\s+дана$"),          # номер экземпляра
    re.compile(r"^\d{4}\s+жылғы"),            # дата утверждения
    re.compile(r"^[А-ЯЁӘҒҚҢӨҰҮҺІ]\.\s?[А-ЯЁӘҒҚҢӨҰҮҺІ][а-яёәғқңөұүһі]+$"),  # «Ш. Жакипов»
    # Строка подписи: звание и через табуляции фамилия с инициалом
    # («полковник \t\t М. Турмагамбетов»).
    re.compile(r"^(?:%s)\s*[\t ]+.*[А-ЯЁӘҒҚҢӨҰҮҺІ]\.\s?[А-ЯЁӘҒҚҢӨҰҮҺІ]" % "|".join(RANKS)),
    # Строка исполнителя («Орынд. Оманов Ж.А.»).
    re.compile(r"^Орынд\.?\s"),
]

#: ПОЗЫВНОЙ — ДВЕ РАЗНЫЕ ЗАПИСИ, И ПЕРВАЯ РЕДАКЦИЯ ВИДЕЛА ТОЛЬКО ОДНУ.
#: Образец пишет и «poz31-44» (число, дефис, число), и «SR-133» (буквы, дефис,
#: число). Прежний образец требовал цифру ПЕРЕД дефисом — и форму «SR-133» не
#: находил вовсе. Из-за этого в бланке остались 73 позывных и 64 настоящие
#: фамилии сотрудников заказчика: «Күзет офицері: Абдрахманов SR-133;».
#: Поймано выгрузкой на живом стенде 27.08.2026 при закрытии Plane №164 —
#: проверка «мест подстановки не осталось» такую строку видеть и не могла,
#: она смотрит на `{{…}}`, а не на людей.
CALLSIGN = r"(?:poz|SR|ПОЗ)\s*[-–]?\s*\d+(?:\s*[-–]\s*\d+)?"

#: Одинокий позывной без фамилии рядом («poz1 poz10-519»): исходный образец
#: ставит их и так. Свой образец нужен потому, что `PERSON` требует фамилию
#: ПЕРЕД позывным и такую запись не видит.
LONE_CALL = re.compile(
    r"(?<![А-Яа-яЁёӘәҒғҚқҢңӨөҰұҮүҺһІі])(%s)" % CALLSIGN, re.I
)

#: Даты мероприятия. В бланке их быть не должно: при следующем событии он
#: напечатал бы чужие числа. Тот же случай, что с датами в «Сводных данных».
DATE = re.compile(r"\b\d{2}(?:-\d{2})?\.\d{2}\.\d{4}\b")
#: Запись человека: фамилия и позывной (`SR-133`, `poz31-44`).
PERSON = re.compile(
    r"([А-ЯЁӘҒҚҢӨҰҮҺІ][А-Яа-яЁёӘәҒғҚқҢңӨөҰұҮүҺһІі\-']+)\s*"
    r"(%s)" % CALLSIGN, re.I
)

#: Фамилия БЕЗ позывного рядом. Нужна потому, что образец пишет людей и так —
#: столбцом в таблице расчёта, где позывной стоит в соседней ячейке. Образец
#: нарочно требует характерное окончание фамилии, а не «слово с большой
#: буквы»: иначе под него попали бы казахские подписи формы («Орналастыруы»,
#: «Бекітемін»), а бланк без подписей — не бланк.
SURNAME = re.compile(
    r"\b[А-ЯЁӘҒҚҢӨҰҮҺІ][а-яёәғқңөұүһі]{2,}"
    r"(?:ов|ев|ин|нов|баев|аев|иев|бек|улы|ұлы|қызы|кызы)\b"
)

def drop(text):
    t = text.strip()
    if t in DROP_LINES:
        return True
    return any(p.match(t) for p in DROP_PATTERNS)

#: Личные поля в СВОЙСТВАХ файла. Их не видно в тексте документа, но Word
#: показывает их в «Свойствах», а при выгрузке они уехали бы заказчику обратно
#: как автор документа системы. В образцах там настоящие ФИО сотрудников.
#: Найдено соседней сессией (Plane №177) — сторож обезличивания читал
#: `word/*.xml` и не заглядывал в `docProps/`.
PERSONAL_PROPS = ("dc:creator", "cp:lastModifiedBy", "Company", "Manager")


def scrub_props(xml_bytes):
    text = xml_bytes.decode("utf-8", "ignore")
    for tag in PERSONAL_PROPS:
        text = re.sub(
            r"<%s>.*?</%s>" % (re.escape(tag), re.escape(tag)),
            "<%s></%s>" % (tag, tag),
            text,
            flags=re.S,
        )
    return text.encode("utf-8")


src = zipfile.ZipFile(SRC)
media = [n for n in src.namelist() if n.startswith("word/media/")]
with zipfile.ZipFile(DST, "w", zipfile.ZIP_DEFLATED) as out:
    for item in src.infolist():
        data = src.read(item.filename)
        if item.filename.startswith("docProps/"):
            data = scrub_props(data)
        if item.filename in media:
            try:
                w, h = Image.open(io.BytesIO(data)).size
                img = Image.new("RGB", (w, h), (238, 240, 244))
                ImageDraw.Draw(img).rectangle([0,0,w-1,h-1], outline=(160,165,175), width=2)
                buf = io.BytesIO(); img.save(buf, "JPEG", quality=80); data = buf.getvalue()
            except Exception:
                pass
        out.writestr(item, data)
src.close()

d = Document(DST)
counter = 0
date_counter = 0
dropped = 0

def process(paragraph):
    global counter, dropped
    text = paragraph.text
    if not text.strip():
        return
    if drop(text):
        for r in paragraph.runs:
            r.text = ""
        dropped += 1
        return
    if not (
        PERSON.search(text)
        or LONE_CALL.search(text)
        or SURNAME.search(text)
        or DATE.search(text)
    ):
        return

    def sub(_m):
        global counter
        counter += 1
        return "{{person_%d}}" % counter

    # ПОРЯДОК ЗАМЕН ЗНАЧИМ. `PERSON` берёт фамилию ВМЕСТЕ с позывным одним
    # местом подстановки; пройди `SURNAME` первым — фамилия ушла бы в своё
    # место, а позывной остался бы сиротой рядом с ним.
    filled = PERSON.sub(sub, text)
    filled = LONE_CALL.sub(sub, filled)
    filled = SURNAME.sub(sub, filled)

    def sub_date(_m):
        global date_counter
        date_counter += 1
        return "{{day_%d}}" % date_counter

    filled = DATE.sub(sub_date, filled)
    if not paragraph.runs:
        return
    paragraph.runs[0].text = filled
    for r in paragraph.runs[1:]:
        r.text = ""

for p in d.paragraphs:
    process(p)
# ОТСЕВА ОБЪЕДИНЁННЫХ ЯЧЕЕК ЗДЕСЬ НЕТ, И ЭТО ВАЖНО. Первая редакция отсеивала
# их по `id(cell._tc)` — и молча пропустила 21 запись. Причина: python-docx
# создаёт обёртки ячеек НА ЛЕТУ, сборщик мусора их освобождает, и `id()`
# переиспользуется. Множество «уже видели» начинало считать виденной ячейку,
# которой не касались ни разу.
#
# Повторная обработка объединённой ячейки безвредна: во второй раз запись уже
# заменена, и образец её не находит. Идемпотентность дешевле хитрого отсева.
for t in d.tables:
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                process(p)
d.save(DST)
print("мест подстановки: людей", counter, "| дат", date_counter, "| снято строк грифа и утверждения:", dropped)

d2 = Document(DST)


def REMAINS(text):
    """Что считается НЕ вычищенным при проверке после сборки.

    Проверять только `PERSON` было недостаточно и оказалось прямой причиной
    утечки: фамилия без позывного и позывной вида «SR-133» под него не
    попадали, а отчёт печатал «осталось записей людей: 0».
    """
    return bool(PERSON.search(text) or LONE_CALL.search(text) or SURNAME.search(text))


left = set()
for p in d2.paragraphs:
    if REMAINS(p.text): left.add(p.text.strip()[:50])
for t in d2.tables:
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                if REMAINS(p.text): left.add(p.text.strip()[:50])
left = sorted(left)
print("осталось записей людей:", len(left))
for x in left[:5]: print("   ", x)
import zipfile as z2
xml = z2.ZipFile(DST).read("word/document.xml").decode("utf-8","ignore")
for word in ("Құпия", "БЕКІТЕМІН", "Жакипов"):
    print(f"«{word}» в бланке:", word in xml)


props = z2.ZipFile(DST)
for name in ("docProps/core.xml", "docProps/app.xml"):
    if name in props.namelist():
        body = props.read(name).decode("utf-8", "ignore")
        for tag in PERSONAL_PROPS:
            m = re.search(r"<%s>(.*?)</%s>" % (re.escape(tag), re.escape(tag)), body, re.S)
            if m and m.group(1).strip():
                print("ОСТАЛОСЬ В СВОЙСТВАХ:", name, tag, "=", m.group(1)[:40])
print("свойства файла проверены")
