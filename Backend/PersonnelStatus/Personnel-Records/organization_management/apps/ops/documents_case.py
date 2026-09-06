"""«Скачать дело» — один документ со всеми вложениями (`[ЗАК-11]`, Plane №437,
Ш-21 плана P2) и «Лист ознакомления» как приложение к расстановке (`[ОЗН-07]`).

Дело ОБЪЕКТА посещения собирается из того, что уже есть в системе, одним
файлом: расстановка сил (текущая + все версии документа из снимков `[СОГ-04]`),
лист ознакомления (ФИО · пост · дата-время · способ), замечания согласования,
оценки этапа 5, журнал штаба (инциденты, замены). Ничего нового дело не
считает — это последний шаг, он только собирает.

DOCX строится с нуля `python-docx` (у дела нет бланка заказчика — это
внутренний архив «для проверок и правоохранительных органов»), PDF — тем же
`emit`, что и у остальных документов.
"""
import datetime as dt
import os
import tempfile

from organization_management.apps.operations.clock import Clock
from organization_management.apps.operations.exceptions import DomainError
from organization_management.apps.operations.models_event import OpsSecurityEvent
from organization_management.apps.ops.documents import emit

#: Право, которым охраняются оценки этапа «Проведение» (Plane №695).
#:
#: То же, что у ручки `visit_object_evaluations`, отдающей ровно эти строки:
#: балл и комментарий по каждому сотруднику. Раздел объявляет своим правилом
#: «выгрузка открывает ровно то, что показывают экраны мероприятия» — и до
#: этой правки нарушал его сам: файл собирался под `event.view`, а те же
#: сведения на экране требовали `event.manage`. Роли EVENT_APPROVER,
#: PATROL_LEAD, DUTY_PLANNER, AUDITOR скачивали закрытое им.
EVALUATIONS_PERMISSION = "event.manage"

_STAGE_ORDER = [
    "BULLETIN", "RECON", "DEMAND", "FORCES", "PLACEMENT", "APPROVAL",
    "ACKNOWLEDGEMENT", "CONDUCT", "CLOSED",
]


def _fmt_dt(value):
    """Отметка времени в зоне РАЗДЕЛА (Plane №696).

    Голый `.astimezone()` брал зону ОПЕРАЦИОННОЙ СИСТЕМЫ, а не `TIME_ZONE`
    проекта: на сервере в UTC каждое время в деле и в приложении к
    расстановке было смещено на пять часов. Локально это невидимо — машина
    разработки уже +05, — поэтому зону берёт `Clock.to_local`, как и весь
    остальной раздел.
    """
    if not value:
        return "—"
    try:
        parsed = dt.datetime.fromisoformat(str(value).replace("Z", "+00:00"))
    except ValueError:
        return str(value)
    return Clock.to_local(parsed).strftime("%d.%m.%Y %H:%M")


def acknowledgement_sheet_rows(event, visit=None):
    """`[ОЗН-07]`: ФИО · пост · дата-время · способ (в системе / лично)."""
    from organization_management.apps.ops import security_events

    posts = security_events.visit_object_posts(event, visit) if visit is not None else (event.recon_sector_posts or [])
    names = {str(p.get("id")): p.get("post") or "" for p in posts}
    rows = []
    for a in event.placement_assignments or []:
        if str(a.get("postId")) not in names:
            continue
        at = a.get("acknowledgedAt")
        manual = a.get("acknowledgedVia") == "personal" or bool(a.get("acknowledgedBy"))
        rows.append([
            a.get("employeeName") or "",
            names.get(str(a.get("postId")), ""),
            _fmt_dt(at) if at else "не подтвердил",
            ("лично" if manual else "в системе") if at else "—",
        ])
    return rows


def _table(document, header, rows, empty="— записей нет —"):
    if not rows:
        document.add_paragraph(empty)
        return
    table = document.add_table(rows=1, cols=len(header))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass  # у бланка заказчика стиля может не быть — сетка не обязательна
    for i, name in enumerate(header):
        table.rows[0].cells[i].text = str(name)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)


def _placement_table(document, event, visit):
    from organization_management.apps.ops.documents_placement import placement_rows

    rows = placement_rows(event, visit)
    _table(
        document,
        ["Сектор", "Пост", "Задача", "Смена", "Потребность", "Назначены"],
        [[r["sector"], r["post"], r["task"], r["shift"], r["need"], r["assigned"]] for r in rows],
        empty="Постов расчёта нет.",
    )


def _versions(document, event, visit):
    from organization_management.apps.operations.models_event import OpsPlacementDocumentVersion

    versions = list(OpsPlacementDocumentVersion.objects.filter(visit_object=visit).order_by("number"))
    if not versions:
        document.add_paragraph("Документ «Расстановка сил» на согласование не отправлялся.")
        return
    for v in versions:
        status = {"DRAFT": "черновик", "SUBMITTED": "на согласовании", "APPROVED": "согласована", "RETURNED": "возвращена"}.get(v.status, v.status)
        document.add_paragraph(
            f"Версия {v.number} — {status}; отправлена {_fmt_dt(v.sent_at.isoformat() if v.sent_at else None)}, "
            f"решение {_fmt_dt(v.decided_at.isoformat() if v.decided_at else None)}"
            + (f", отменена {_fmt_dt(v.superseded_at.isoformat())}" if v.superseded_at else "")
        )
        snapshot = v.snapshot or {}
        names = {str(p.get("id")): p.get("post") or "" for p in snapshot.get("posts") or []}
        _table(
            document,
            ["Пост", "Сотрудник", "Роль"],
            [[names.get(str(a.get("postId")), a.get("postId")), a.get("employeeName") or "", a.get("roleCode") or ""]
             for a in snapshot.get("assignments") or []],
            empty="В снимке версии назначений нет.",
        )


def _post_labels(event):
    """`{id поста → «сектор · пост»}` по расчёту мероприятия.

    Считается ОДИН раз на объект, а не на замечание: расчёт лежит полем
    строки, и перебирать его в цикле по замечаниям значило бы делать ту же
    работу столько раз, сколько замечаний.
    """
    from organization_management.apps.ops.security_events import post_label

    return {
        str(post.get("id")): post_label(post)
        for post in event.recon_sector_posts or []
        if post.get("id")
    }


def _remark_attachment(remark, labels):
    """Привязка замечания для колонки «Привязка» дела.

    🔴 СНЯТЫЙ ПОСТ НАЗЫВАЕТСЯ, А НЕ ПРЕВРАЩАЕТСЯ В «ОБЩЕЕ» (Plane №510). При
    снятии поста замечание отвязывается от него (иначе оно ссылалось бы на
    несуществующий пост и невидимо держало бы согласование), но согласующий
    писал про КОНКРЕТНЫЙ пост — и в деле, которое читают потом, «общее»
    сказало бы неправду.

    🔴 ПОСТ НАЗЫВАЕТСЯ ПОДПИСЬЮ, А НЕ ИДЕНТИФИКАТОРОМ (Plane №865). Здесь
    стояло `remark.get("postName") or remark.get("postId")`, а ключ
    `postName` замечанию не пишет НИКТО: `new_remark` его не заводит. То есть
    ветка была мертва, и колонка «Привязка» подписываемого и рассылаемого
    дела печатала `post-4f2a9c1b8e30` у каждого привязанного замечания.
    Экранная половина при этом была в порядке — `ApprovalStage` разрешает id
    в «сектор · пост», — и расхождение видел только тот, кто открывал файл.

    `postName` оставлен первым по счёту намеренно: если замечание однажды
    начнут писать с готовой подписью, она победит расчёт, а не наоборот.
    Идентификатор остаётся последним запасом — на пост, которого в расчёте
    уже нет, а отвязать его забыли.
    """
    detached = str(remark.get("detachedPost") or "")
    post_id = remark.get("postId")
    named = remark.get("postName") or (labels.get(str(post_id)) if post_id else None) or post_id
    if named:
        return str(named)
    if detached:
        return f"{detached} (пост снят с расчёта)"
    return "общее"


def _remarks(document, visit, event):
    labels = _post_labels(event)
    rows = []
    for r in visit.approval_remarks or []:
        status = {"OPEN": "открыто", "RESOLVED": "устранено", "DISAGREED": "не согласен"}.get(r.get("status"), r.get("status") or "—")
        rows.append([
            r.get("text") or "",
            r.get("authorName") or r.get("author") or "",
            _fmt_dt(r.get("createdAt")),
            _remark_attachment(r, labels),
            ("срочно, " if r.get("urgent") else "") + status,
            (r.get("response") or "") + (f" ({_fmt_dt(r.get('respondedAt'))})" if r.get("respondedAt") else ""),
            # Замечания, поставленные до версий документа (№398), номера не несут.
            (f"v{r.get('documentVersion')}" if r.get("documentVersion") else "—")
            + (f" → v{r.get('resolvedInDocumentVersion')}" if r.get("resolvedInDocumentVersion") else ""),
        ])
    _table(document, ["Замечание", "Автор", "Дата", "Привязка", "Статус", "Ответ старшего", "Версия"], rows, empty="Замечаний не было.")


def _may_see_evaluations(permissions):
    """`None` — прав не проверяем (внутренний вызов), иначе обычная мерка.

    Умолчание НЕ «запретить»: сборщик зовут не только ручка, и собственного
    понятия о спрашивающем у него нет. Запрет по умолчанию означал бы, что
    внутренний вызов молча теряет раздел вместо отказа.
    """
    if permissions is None:
        return True
    return "*" in permissions or EVALUATIONS_PERMISSION in permissions


def _evaluations(document, event, visit, permissions=None):
    from organization_management.apps.ops import conduct_evaluations

    if not _may_see_evaluations(permissions):
        # ИЗЪЯТИЕ НАЗЫВАЕТСЯ ВСЛУХ, а раздел остаётся в оглавлении. Дело —
        # архивный документ «для проверок»: молча собранный без раздела, он
        # выглядел бы полным, и два человека получили бы разные файлы под
        # одним именем, не зная об этом.
        document.add_paragraph(
            "Раздел не показан: у вас нет права "
            f"«{EVALUATIONS_PERMISSION}». Оценки и комментарии по сотрудникам "
            "открыты тому, кто ведёт мероприятие."
        )
        return
    summary = conduct_evaluations.visit_evaluations(event, visit)
    rows = [[r["sector"], r["post"], r["employeeName"] + (" (снят)" if r["replaced"] else ""), r["score"] if r["score"] is not None else "—", r["comment"]]
            for r in summary["rows"]]
    document.add_paragraph(f"Оценено {summary['evaluated']} из {summary['total']}.")
    _table(document, ["Сектор", "Пост", "Сотрудник", "Оценка", "Комментарий"], rows, empty="Оценивать было некого.")


def _journal(document, event):
    kinds = {"INCIDENT": "инцидент", "REPLACEMENT": "замена", "NOTE": "запись"}
    rows = [[_fmt_dt(e.get("createdAt")), kinds.get(e.get("type"), e.get("type") or ""), e.get("title") or "", e.get("description") or ""]
            for e in event.journal_entries or []]
    _table(document, ["Время", "Вид", "Заголовок", "Описание"], rows, empty="Инцидентов не было.")


def render_case(event_code, *, visit_object_id=None, fmt="pdf", permissions=None):
    """Байты дела объекта (все объекты ОМ, если объект не назван).

    `permissions` — права спрашивающего (Plane №695). `None` означает
    «прав не проверяем»: сборщик зовут и внутренние пути, у которых
    спрашивающего нет вовсе.
    """
    from docx import Document
    from docx.enum.text import WD_BREAK

    event = OpsSecurityEvent.objects.filter(code=event_code).first()
    if event is None:
        raise DomainError("ENTITY_NOT_FOUND", 404, detail={"code": [str(event_code)]}, message="Мероприятие не найдено.")
    # Объект не назван — дело ВСЕГО мероприятия, по объекту на раздел (в
    # отличие от расстановки, где без объекта при нескольких — отказ).
    visits = list(event.visit_objects.order_by("position", "pk"))
    if visit_object_id not in (None, ""):
        visits = [v for v in visits if str(v.pk) == str(visit_object_id)]
        if not visits:
            raise DomainError("ENTITY_NOT_FOUND", 404, detail={"id": str(visit_object_id)}, message="Объект посещения не найден.")
    if not visits:
        # ОМ без объектов посещения (заведено до №385): дело по мероприятию.
        visits = [None]
    document = Document()
    document.add_heading(f"Дело · {event.code} — {event.title}", level=0)
    document.add_paragraph(
        f"Дата мероприятия {event.business_date.strftime('%d.%m.%Y')} · собрано {Clock.local_now().strftime('%d.%m.%Y %H:%M')}"
        + (f" · мероприятие закрыто {_fmt_dt(event.closed_at.isoformat())}" if event.closed_at else "")
    )
    for index, visit in enumerate(visits):
        if index:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        if visit is None:
            document.add_heading(f"Объект «{event.object_name or '—'}»", level=1)
        else:
            document.add_heading(f"Объект «{visit.object_name}»", level=1)
            document.add_paragraph(
                f"Старший объекта: {visit.chief_name or 'не назначен'} · состояние: {visit.stage}"
                + (f" · закрыт {_fmt_dt(visit.closed_at.isoformat())}" if visit.closed_at else "")
                + (f" · итог: {visit.closing_comment}" if visit.closing_comment else "")
            )
        document.add_heading("1. Расстановка сил", level=2)
        _placement_table(document, event, visit)
        document.add_heading("1.1. Версии документа «Расстановка сил»", level=3)
        if visit is None:
            document.add_paragraph("У мероприятия без объектов посещения версий документа нет.")
        else:
            _versions(document, event, visit)
        document.add_heading("2. Лист ознакомления", level=2)
        _table(document, ["ФИО", "Пост", "Дата-время", "Способ"], acknowledgement_sheet_rows(event, visit), empty="Назначений не было.")
        document.add_heading("3. Замечания согласования", level=2)
        if visit is None:
            document.add_paragraph("Замечаний не было.")
        else:
            _remarks(document, visit, event)
        document.add_heading("4. Оценки сотрудников", level=2)
        if visit is None:
            document.add_paragraph("Оценивать было некого.")
        else:
            _evaluations(document, event, visit, permissions)
    document.add_heading("5. Журнал штаба", level=1)
    _journal(document, event)
    handle, path = tempfile.mkstemp(suffix=".docx", prefix="delo-")
    os.close(handle)
    try:
        document.save(path)
        return emit(path, fmt)
    finally:
        try:
            os.remove(path)
        except OSError:
            pass


def append_acknowledgement_sheet(document, event, visit):
    """`[ОЗН-07]`: приложение к расстановке после завершения ознакомления."""
    from docx.enum.text import WD_BREAK

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    # Не `add_heading`: у бланка заказчика нет стилей Heading, и python-docx
    # падает на «no style with name».
    document.add_paragraph().add_run("Приложение. Лист ознакомления").bold = True
    _table(document, ["ФИО", "Пост", "Дата-время", "Способ"], acknowledgement_sheet_rows(event, visit), empty="Назначений не было.")


def acknowledgement_completed(event, visit=None):
    """Ознакомление пройдено — можно прикладывать лист (`[МД-08]`).

    🔴 СПРАШИВАЕТСЯ ЭТАП ОБЪЕКТА, КОГДА ОБЪЕКТ НАЗВАН (Plane №520). Функция
    сравнивала только `event.stage`, а он — НАИМЕНЬШАЯ стадия среди объектов.
    Документ же печатается ПО ОБЪЕКТУ: объект А ознакомление закончил, объект Б
    отстаёт — и в документе объекта А приложения не было, потому что
    мероприятие целиком ещё на «Ознакомлении». Человек скачивал дело
    закончившего объекта и не находил в нём листа, который этот объект как раз
    и подписал.

    Это третий по счёту случай одной и той же ошибки — «спросили мероприятие
    там, где отвечает объект» (см. №475 и №528). Общее у них одно: с №412
    `event.stage` перестал быть ответом на вопрос «где этот объект».

    Объект НЕ НАЗВАН (ОМ без объектов посещения, сборка по всему мероприятию) —
    отвечает мероприятие, как и раньше: другой сущности у такого ОМ нет.
    """
    stage = visit.stage if visit is not None else event.stage
    return (
        stage in _STAGE_ORDER
        and _STAGE_ORDER.index(stage) >= _STAGE_ORDER.index("CONDUCT")
    )
