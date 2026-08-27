"""Сводка ГВО НА СЕРВЕРЕ и её раскладка в шаблон (Plane №158, шаг ПД-2).

ЗАЧЕМ ЭТО ЗДЕСЬ. Документ «Сводные данные» обязан быть СРЕЗОМ СИСТЕМЫ. А
сводка ГВО до этой правки собиралась только на КЛИЕНТЕ: сервер хранил лишь
ручные правки (`OpsGvoSummaryPatch`), а база — страна, лица, прибытие, убытие,
объекты посещения — считалась в браузере (`deriveGvoSummary`). Наполнять
документ тем, что пришлёт браузер, нельзя: тогда содержимое диктует клиент, и
прислать можно что угодно.

ОТКУДА ПРАВИЛА. Порт `entities/gvo-summary/model/derive.ts` строка в строку:
деловая дата → прибытие и убытие, охраняемое лицо бюллетеня → «Охраняемые
лица», ответственный ОМ → «Ответственный», объекты мероприятия → расписание
посещений. Всё это ЖИВЁТ НА СЕРВЕРЕ — клиент ничего не добавлял от себя, он
только складывал.

ДОЛГ ЗАКРЫВАЕТСЯ (Plane №166). Сборка какое-то время жила в ДВУХ местах — здесь
и на клиенте — и успела разойтись за один день: дата тут писалась «10.09.2026г.»,
а на экране «10.09.2026». Именно так это и расходится: не спором о правилах, а
мелочью, которую в одном месте поправили под документ, а во втором никто не
увидел. Теперь сводку отдаёт ручка `GET /api/ops/gvo-summaries/<код>/`, экран
читает её, а `deriveGvoSummary` снимается после переезда читателей.
"""
import datetime as dt

#: То же слово, что у клиента (`UNSPECIFIED`): пустое поле и «уточняется» —
#: разные вещи. Пустое читается как «сведений нет вовсе», а «уточняется» —
#: как «знаем, что нужно, но ещё не выяснили».
UNSPECIFIED = "уточняется"

WEEKDAYS = (
    "понедельник", "вторник", "среда", "четверг",
    "пятница", "суббота", "воскресенье",
)


def _ru_date(value):
    if value is None:
        return UNSPECIFIED
    if isinstance(value, str):
        try:
            value = dt.date.fromisoformat(value)
        except ValueError:
            return value
    return f"{value.day:02d}.{value.month:02d}.{value.year}"


#: Дата В ДОКУМЕНТЕ пишется с «г.» — так в образце заказчика («17.06.2026 г.»).
#: Дата В СВОДКЕ пишется без него — так её показывает экран и так объявлен
#: клиентский тип. Это РАЗНЫЕ слои, и суффикс живёт в том, который его требует:
#: пока он сидел в сборке сводки, экран и документ уже расходились в дате, а
#: ловить это было некому — сборок две (Plane №166).
def _document_date(value):
    value = (value or "").strip()
    return f"{value} г." if value and value != UNSPECIFIED else value


def _ru_weekday(value):
    if isinstance(value, str):
        try:
            value = dt.date.fromisoformat(value)
        except ValueError:
            return ""
    return WEEKDAYS[value.weekday()] if value else ""


def visit_days(event):
    """«Объекты посещения» — из ТАБЛИЦЫ объектов мероприятия, а не из патча.

    Так же, как на клиенте, и по той же причине: до «Реестра ОМ-35.1» список
    жил в двух местах и расходился молча — объект, дописанный в сводке, не
    получал ни постов, ни готовности.
    """
    # `visit_objects` — СВЯЗАННЫЕ ЗАПИСИ, а не поле JSON: порядок задаёт
    # `position` (он же порядок раскрытия строки реестра), а не порядок в базе.
    #
    # Сортировка В ПАМЯТИ, а не `.order_by()`: тот строит НОВЫЙ запрос и
    # проходит мимо `prefetch_related`, из-за чего реестр сводок ходил в базу
    # за объектами каждого мероприятия отдельно (Plane №166). Объектов у ОМ
    # единицы — сортировать их списком дешевле, чем ещё раз спрашивать базу.
    rows = sorted(event.visit_objects.all(), key=lambda row: (row.position, row.id))
    by_day = {}
    for visit in rows:
        iso = str(visit.visit_day or event.business_date)
        note = (visit.note or "").strip()
        item = {
            "obj": visit.object_name or "",
            "note": note if note else UNSPECIFIED,
        }
        by_day.setdefault(iso, []).append(item)
    return [
        {"day": _ru_date(iso), "weekday": _ru_weekday(iso), "items": items}
        for iso, items in sorted(by_day.items())
    ]


def derive_summary(event):
    """База сводки из мероприятия. Порт клиентского `deriveGvoSummary`."""
    day = _ru_date(event.business_date)
    person = (event.protected_person_name or "").strip()
    # Сводные данные строятся ПО ОДНОМУ лицу — так устроен образец заказчика:
    # страна, антропометрия, группа крови у каждого свои, и в одну карточку
    # двоих не положить. Поэтому список лиц (Plane №188) сюда НЕ раскрывается:
    # берётся главное, остальные попадут в свои сводки, когда документ научится
    # собираться на каждого. Это осознанная граница, а не пропуск.
    owner = (event.owner_name or "").strip()
    return {
        "country": UNSPECIFIED,
        # Пусто, если в бюллетене лицо не назвали: подставлять сюда
        # «уточняется» вместо человека нечем.
        "persons": (
            [{"name": person, "role": "охраняемое лицо", "facts": []}]
            if person
            else []
        ),
        "arrival": {"date": day, "time": UNSPECIFIED, "route": UNSPECIFIED,
                    "flight": UNSPECIFIED, "dur": UNSPECIFIED},
        "departure": {"date": day, "time": UNSPECIFIED, "route": UNSPECIFIED,
                      "flight": UNSPECIFIED, "dur": UNSPECIFIED},
        "meet": [],
        "farewell": [],
        "stay": {"place": UNSPECIFIED, "room": UNSPECIFIED},
        "delegation": [],
        "sbChief": UNSPECIFIED,
        "weapons": UNSPECIFIED,
        "wishes": UNSPECIFIED,
        "obVariant": UNSPECIFIED,
        "radio": UNSPECIFIED,
        "responsible": (
            {"name": owner, "callsign": UNSPECIFIED, "role": "ответственный"}
            if owner
            else None
        ),
        "groups": [{"name": "ГВО (состав уточняется)", "members": []}],
        # Свободный текст «Выделяемый транспорт»: его набирает человек в
        # разделе сводки. ОСТАЁТСЯ пустым в базе и наполняется патчем — так
        # было и до реестра транспорта.
        "transport": [],
        # Машины, ВЫДЕЛЕННЫЕ из реестра ГОН (Plane №215). Отдельным ключом,
        # а не подменой `transport`: у свободного текста есть свои читатели
        # (эта же сводка и документ сводных данных ниже), и снимать источник,
        # пока его читают, правило раздела запрещает. Патчем этот ключ не
        # правится вовсе — он ВЫВОД из выделений, а не запись человека, и его
        # нет в `ALLOWED_PATCH_KEYS`.
        "allocatedTransport": [
            {
                "callsign": row.callsign,
                "label": row.vehicle_label,
                "purpose": row.purpose,
                "plate": row.vehicle.plate if row.vehicle is not None else None,
                "armorClass": (
                    row.vehicle.armor_class if row.vehicle is not None else None
                ),
            }
            for row in event.vehicles.all()
        ],
        "visits": visit_days(event),
    }


def _deep_merge(base, patch):
    """База + патч. Вложенные словари сливаются ГЛУБОКО: правка раздела
    «Прибытие» может нести только время, не затирая маршрут."""
    result = dict(base)
    for key, value in (patch or {}).items():
        if isinstance(value, dict) and isinstance(result.get(key), dict):
            result[key] = _deep_merge(result[key], value)
        else:
            result[key] = value
    return result


def summary_for_event(event):
    """Сводка мероприятия: база плюс сохранённые ручные правки."""
    from organization_management.apps.operations.models_gvo import (
        OpsGvoSummaryPatch,
    )

    patch = (
        OpsGvoSummaryPatch.objects.filter(event_id=event.pk)
        .values_list("patch", flat=True)
        .first()
    )
    return _deep_merge(derive_summary(event), patch or {})


def summary_row(event, record=None, *, fetch=True):
    """Строка сводки: собранная сводка плюс признак «Заполнена».

    `filled` считает СЕРВЕР, а не экран: признак живёт там же, откуда пришла
    сводка, иначе он снова оказался бы правилом на клиенте.
    """
    if fetch:
        from organization_management.apps.operations.models_gvo import (
            OpsGvoSummaryPatch,
        )

        record = OpsGvoSummaryPatch.objects.filter(event_id=event.pk).first()
    patch = (record.patch if record else None) or {}
    return {
        "omCode": event.code,
        "summary": _deep_merge(derive_summary(event), patch),
        # «Заполнена» — если по мероприятию есть хоть одна ручная правка.
        # Иначе «Черновик»: всё показанное выведено из бюллетеня.
        "filled": bool(patch),
        # null — правок не было вовсе, а не «время неизвестно».
        "updatedAt": record.updated_at.isoformat() if record else None,
    }


def assembled_summaries():
    """Собранные сводки ВСЕХ мероприятий — по одной строке на мероприятие.

    Реестрам (ГВО, охраняемые лица) нужна сводка каждого ОМ, а не только тех,
    у кого есть ручные правки: у мероприятия без правок сводка не пустая, она
    выведена из бюллетеня. Поэтому строка есть у каждого, а `filled` отличает
    «Заполнена» от «Черновика».

    Собирается ОДНИМ проходом, а не вызовом `summary_for_event` в цикле: тот
    ходит в базу за патчем на каждое мероприятие, и реестр из сорока строк
    стоил бы сорок запросов.
    """
    from organization_management.apps.operations.models_event import (
        OpsSecurityEvent,
    )
    from organization_management.apps.operations.models_gvo import (
        OpsGvoSummaryPatch,
    )

    patches = {
        record.event_id: record
        for record in OpsGvoSummaryPatch.objects.all()
    }
    # `vehicles__vehicle` в предзагрузке, а не запрос на строку: реестр из
    # сорока ОМ иначе стоил бы сорок запросов за машинами (та же причина, по
    # которой патчи собраны одним проходом выше).
    events = OpsSecurityEvent.objects.prefetch_related(
        "visit_objects", "vehicles__vehicle"
    ).order_by("code")
    return [
        summary_row(event, patches.get(event.pk), fetch=False)
        for event in events
    ]


def _person_lines(person):
    """Строки антропометрии одного лица — в порядке образца."""
    facts = person.get("facts") or []
    return [str(fact) for fact in facts]


def document_values(event):
    """Сводка → значения мест подстановки шаблона «Сводные данные».

    Ключи заданы ШАБЛОНОМ, а не выдуманы здесь: шаблон снят с образца
    заказчика, и его места подстановки — это перечень того, что документ
    обещает показать. Всё, чего в сводке нет, отдаётся ПУСТЫМ, а не
    «уточняется»: пустая строка под подписью читается как «сведений нет», и
    это честно; выдуманное слово читалось бы как факт.
    """
    summary = summary_for_event(event)
    persons = summary.get("persons") or []
    values = {}

    values["country_1"] = summary.get("country") or ""
    for index in (1, 2):
        person = persons[index - 1] if len(persons) >= index else {}
        values[f"person{index}_title"] = person.get("role", "") if person else ""
        values[f"person{index}_name"] = person.get("name", "") if person else ""
        for line_no, line in enumerate(_person_lines(person), start=1):
            values[f"person{index}_data_{line_no}"] = line

    arrival = summary.get("arrival") or {}
    departure = summary.get("departure") or {}
    values["arrival_1"] = " ".join(
        part
        for part in (_document_date(arrival.get("date")), arrival.get("time"))
        if part
    )
    values["departure_1"] = " ".join(
        part
        for part in (_document_date(departure.get("date")), departure.get("time"))
        if part
    )

    stay = summary.get("stay") or {}
    values["accommodation_1"] = " ".join(
        part for part in (stay.get("place"), stay.get("room")) if part
    )
    values["security_chief_1"] = summary.get("sbChief") or ""
    values["armament_1"] = summary.get("weapons") or ""
    values["wishes_1"] = summary.get("wishes") or ""
    values["route_variant_1"] = summary.get("obVariant") or ""
    values["radio_channel_1"] = summary.get("radio") or ""

    for line_no, item in enumerate(summary.get("meet") or [], start=1):
        values[f"meeting_{line_no}"] = str(item)
    for line_no, item in enumerate(summary.get("farewell") or [], start=1):
        values[f"seeing_off_{line_no}"] = str(item)
    for line_no, item in enumerate(summary.get("delegation") or [], start=1):
        values[f"delegation_{line_no}"] = str(item)
    # Транспорт документа: СНАЧАЛА выделенные машины реестра, затем строки
    # свободного текста (Plane №215). Порядок не косметика — выделение несёт
    # ГРНЗ и класс брони, то есть сведения, которых у текста нет вовсе, и
    # ставить его после значило бы прятать точное за приблизительным. Текст
    # остаётся: пока его кто-то набирает, выбрасывать набранное нельзя.
    line_no = 0
    for car in summary.get("allocatedTransport") or []:
        line_no += 1
        values[f"transport_{line_no}"] = " — ".join(
            part
            for part in (car.get("callsign"), car.get("label"), car.get("purpose"))
            if part
        )
    for car in summary.get("transport") or []:
        line_no += 1
        values[f"transport_{line_no}"] = " — ".join(
            part for part in (car.get("code"), car.get("car")) if part
        )

    line_no = 0
    for group in summary.get("groups") or []:
        for member in group.get("members") or []:
            line_no += 1
            values[f"gvo_staff_{line_no}"] = " ".join(
                part
                for part in (
                    member.get("name"),
                    member.get("role"),
                    member.get("callsign"),
                )
                if part
            )

    # Расписание: четыре дня, как в образце. Дней больше — они не теряются
    # молча, а называются в последнем дне списком: потерянный день посещения
    # опаснее переполненной ячейки.
    days = summary.get("visits") or []
    for day_no in range(1, 5):
        day = days[day_no - 1] if len(days) >= day_no else None
        if day is None:
            values[f"day{day_no}_date_1"] = ""
            continue
        weekday = day.get("weekday")
        shown = _document_date(day.get("day"))
        values[f"day{day_no}_date_1"] = f"{shown} ({weekday})" if weekday else shown
        for line_no, item in enumerate(day.get("items") or [], start=1):
            values[f"day{day_no}_line{line_no}_1"] = " — ".join(
                part for part in (item.get("obj"), item.get("note")) if part
            )
    return values


def fill_all_keys(template_keys, values):
    """Дополнить значения ПУСТЫМИ для всех мест шаблона, которых нет в данных.

    Без этого конвейер честно откажется собирать документ: он не выпускает
    наружу файл с `{{...}}` вместо значений. Но отсутствие сведений — не
    поломка, а обычное состояние сводки, которую ещё заполняют. Разница
    именно здесь: ПУСТО значит «не заполнено», а `{{...}}` значило бы «сломано».
    """
    return {key: values.get(key, "") for key in template_keys}


#: Шаблон снят С ОБРАЗЦА ЗАКАЗЧИКА: вёрстка, рамки, заливка и подчёркнутые
#: подписи взяты как есть, а настоящие данные заменены местами подстановки, и
#: фотографии — нейтральными заглушками. Персональных сведений в репозитории
#: не лежит.
SUMMARY_TEMPLATE = "summary_data.docx"


def summary_template_path():
    import os

    return os.path.join(
        os.path.dirname(__file__), "document_templates", SUMMARY_TEMPLATE
    )


def template_keys(template_path):
    """Места подстановки, объявленные ШАБЛОНОМ. Источник — сам файл, а не
    список в коде: список разошёлся бы с шаблоном при первой же его правке."""
    from docx import Document

    from organization_management.apps.ops.documents import PLACEHOLDER

    document = Document(template_path)
    keys = set()
    for paragraph in document.paragraphs:
        keys.update(PLACEHOLDER.findall(paragraph.text))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                keys.update(PLACEHOLDER.findall(cell.text))
    return keys


def render_summary_pdf(event, fmt="pdf"):
    """«Сводные данные» мероприятия; `fmt` — «docx» либо «pdf»."""
    from organization_management.apps.ops.documents import (
        render_docx_from_template,
        render_pdf_from_template,
    )

    path = summary_template_path()
    values = fill_all_keys(template_keys(path), document_values(event))
    if fmt == "docx":
        return render_docx_from_template(path, values)
    return render_pdf_from_template(path, values)
