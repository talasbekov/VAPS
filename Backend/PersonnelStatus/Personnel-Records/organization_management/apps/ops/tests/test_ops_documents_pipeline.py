"""Конвейер «шаблон → подстановка → PDF» (Plane №157, шаг ПД-1).

Заказчик: документы обязаны выглядеть В ТОЧНОСТИ как его файлы Word. Значит
документ не рисуется заново, а берётся готовым — меняются только значения.
Отсюда и предмет проб: не «PDF собрался», а «значения подставились, и ничего
недозаполненного наружу не ушло».

Быстрые пробы работают на уровне `.docx` — там видно, что именно подставилось.
Полная конвертация дорогая (внешний процесс), поэтому она проверяется ОДНОЙ
сквозной пробой: PDF собран и текст из него читается.
"""
import os
import shutil
import subprocess
import tempfile

import pytest
from docx import Document

from organization_management.apps.operations.exceptions import DomainError
from organization_management.apps.ops import documents

TEMPLATE = os.path.join(
    os.path.dirname(documents.__file__),
    "document_templates",
    "pipeline_probe.docx",
)

VALUES = {
    "snapshot_date": "27.08.2026",
    "snapshot_time": "08:00",
    "title": "Сводные данные",
    "country": "Черногория",
    "person": "Яков Милатович",
    "blood_group": "А (II) Rh +",
}


def filled_text(values, template=TEMPLATE):
    """Заполненный документ как текст: тело плюс таблицы."""
    path, left = documents.fill_template(template, values)
    try:
        document = Document(path)
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts), left
    finally:
        os.unlink(path)


def test_values_replace_placeholders(db):
    text, left = filled_text(VALUES)

    assert "Черногория" in text
    assert "Яков Милатович" in text
    assert "А (II) Rh +" in text
    assert left == []
    # Мест подстановки не осталось НИ ОДНОГО — иначе документ уйдёт наружу
    # с «{{country}}» вместо страны и будет выглядеть готовым.
    assert "{{" not in text


def test_placeholder_split_across_runs_is_still_replaced(db):
    """Главная проба шага, и она про то, что ломается МОЛЧА.

    Word режет текст на «прогоны» по границам форматирования и делает это
    непредсказуемо: `{{country}}` легко оказывается разложенным на `{{cou`,
    `ntry`, `}}`. Замена по каждому прогону отдельно не найдёт ничего и НЕ
    сообщит об этом — документ просто выйдет с местами подстановки.

    Здесь разрыв воспроизводится нарочно: место подстановки собирается из трёх
    прогонов, как это делает сам Word.
    """
    document = Document()
    paragraph = document.add_paragraph()
    for piece in ("{{cou", "ntry", "}}"):
        paragraph.add_run(piece)
    handle, path = tempfile.mkstemp(suffix=".docx")
    os.close(handle)
    document.save(path)
    try:
        text, left = filled_text({"country": "Черногория"}, template=path)
    finally:
        os.unlink(path)

    assert "Черногория" in text
    assert left == []


def test_unfilled_document_does_not_leave_the_building(db):
    """Пустое значение и НЕЗАПОЛНЕННОЕ место — разные вещи.

    Пустое поле заказчик прочтёт как «сведений нет»; `{{person}}` он прочтёт
    как поломку, и будет прав. Поэтому недозаполненный документ не
    конвертируется вовсе.
    """
    with pytest.raises(DomainError) as failure:
        documents.render_pdf_from_template(TEMPLATE, {"country": "Черногория"})

    assert failure.value.code == "DOCUMENT_INCOMPLETE"
    assert "person" in failure.value.detail["placeholders"]


def test_missing_template_is_named(db):
    with pytest.raises(DomainError) as failure:
        documents.fill_template("/nowhere/нет-такого.docx", VALUES)

    assert failure.value.code == "DOCUMENT_TEMPLATE_MISSING"


@pytest.mark.skipif(
    shutil.which("soffice") is None or shutil.which("pdftotext") is None,
    reason="нужны soffice и pdftotext — сквозная проба идёт только там, где есть конвертер",
)
def test_pdf_carries_the_values_as_text(db):
    """Сквозная проба: PDF собран, и значения в нём — ТЕКСТ, а не картинка.

    Проверять размер файла бессмысленно: PDF из пустого шаблона тоже не пуст.
    """
    pdf = documents.render_pdf_from_template(TEMPLATE, VALUES)
    handle, path = tempfile.mkstemp(suffix=".pdf")
    os.close(handle)
    try:
        with open(path, "wb") as out:
            out.write(pdf)
        text = subprocess.run(
            ["pdftotext", path, "-"], capture_output=True, timeout=60
        ).stdout.decode("utf-8", "ignore")
    finally:
        os.unlink(path)

    assert pdf[:4] == b"%PDF"
    assert "Черногория" in text
    assert "Яков Милатович" in text
    # Кириллица настоящая, а не квадраты: если бы шрифт потерялся, извлечённый
    # текст был бы пустым или мусорным.
    assert "Сводные данные" in text
    assert "{{" not in text


def test_missing_converter_is_refused_not_substituted(db, monkeypatch):
    """Нет конвертера — отказ, а не подмена формата.

    Отдать `.docx` вместо PDF значило бы решить за заказчика, что ему подойдёт
    другой формат. Он просил PDF.
    """
    monkeypatch.setattr(documents.shutil, "which", lambda _: None)

    with pytest.raises(DomainError) as failure:
        documents.docx_to_pdf("/tmp/что-угодно.docx")

    assert failure.value.code == "PDF_CONVERTER_MISSING"


def test_broken_template_is_named_not_swallowed(db, tmp_path):
    """Битый файл называется файлом и причиной, а не «не работает».

    Поймано на настоящем образце заказчика `01 Сводные данные РЭС 22.04.docx`:
    у него нет конца zip-архива и центрального каталога — файл обрезан.
    Команда `file` при этом отвечает «Microsoft Word 2007+», потому что
    смотрит только на первые байты, а LibreOffice — невнятное «source file
    could not be loaded». Без явного отказа это читается как поломка выгрузки,
    и чинить будут не то.
    """
    broken = tmp_path / "обрезанный.docx"
    # Начало настоящего zip и обрыв: ровно то состояние, что у образца.
    broken.write_bytes(b"PK\x03\x04\x14\x00\x06\x00" + b"\x00" * 64)

    with pytest.raises(DomainError) as failure:
        documents.fill_template(str(broken), VALUES)

    assert failure.value.code == "DOCUMENT_TEMPLATE_BROKEN"
    assert "обрезанный.docx" in failure.value.detail["template"][0]
