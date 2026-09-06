"""«Скачать дело» (`[ЗАК-11]`, Plane №437) и «Лист ознакомления» как
приложение к расстановке (`[ОЗН-07]`).

Проверяется в DOCX (тот же документ до конвертации; PDF делает LibreOffice
и на тестовой машине его может не быть): в деле есть все секции и живые
данные объекта; расстановка после «Ознакомления» несёт приложение, до —
нет; реестр знает вид `case`, ручка отдаёт файл.
"""
import io

import pytest
from docx import Document

from organization_management.apps.ops import documents_case, documents_registry
from organization_management.apps.ops.documents_placement import render_placement
from organization_management.apps.ops.tests.test_ops_visit_object_close import (  # noqa: F401
    actor,
    two_objects_on_conduct,
)
from organization_management.apps.ops.tests.test_ops_visit_object_approval import (  # noqa: F401
    two_objects_on_approval,
)
from organization_management.apps.ops.tests.test_ops_security_events_api import (  # noqa: F401
    URL,
    approver,
    make_employee,
    make_object,
    manager,
    viewer,
)

pytestmark = pytest.mark.django_db


def _text(payload):
    document = Document(io.BytesIO(payload))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_case_carries_every_section_of_the_object(manager, two_objects_on_conduct):  # noqa: F811
    _, event_id, first, _ = two_objects_on_conduct
    from organization_management.apps.ops import security_events as service
    event = service.lock_event(event_id)
    payload = documents_case.render_case(event.code, visit_object_id=str(first.pk), fmt="docx")
    text = _text(payload)
    for section in ["Расстановка сил", "Версии документа", "Лист ознакомления", "Замечания согласования", "Оценки сотрудников", "Журнал штаба"]:
        assert section in text, section
    assert first.object_name in text
    # Живые данные: назначенный на пост объекта и версия документа.
    assigned = [a for a in event.placement_assignments if a.get("employeeName")]
    assert assigned and assigned[0]["employeeName"] in text
    assert "Версия 1" in text


def test_case_without_object_collects_all_objects(manager, two_objects_on_conduct):  # noqa: F811
    _, event_id, first, second = two_objects_on_conduct
    from organization_management.apps.ops import security_events as service
    event = service.lock_event(event_id)
    text = _text(documents_case.render_case(event.code, fmt="docx"))
    assert first.object_name in text and second.object_name in text


def test_placement_gets_the_sheet_only_after_acknowledgement(manager, two_objects_on_conduct):  # noqa: F811
    _, event_id, first, _ = two_objects_on_conduct
    from organization_management.apps.ops import security_events as service
    event = service.lock_event(event_id)
    after = _text(render_placement(event.code, fmt="docx", visit_object_id=str(first.pk)))
    assert "Приложение. Лист ознакомления" in after
    # До «Проведения» приложения нет (🔴 мутация: порог стадии).
    #
    # Отодвигается стадия ОБЪЕКТА, а не мероприятия (Plane №520): документ
    # печатается по объекту, и порог спрашивается у него. Раньше здесь
    # двигался `event.stage` — и проба проходила ровно потому, что порог
    # спрашивали не у того, у кого документ.
    first.stage = "PLACEMENT"
    first.save(update_fields=["stage", "updated_at"])
    before = _text(render_placement(event.code, fmt="docx", visit_object_id=str(first.pk)))
    assert "Приложение. Лист ознакомления" not in before


def test_a_finished_object_keeps_its_sheet_while_the_neighbour_lags(manager, two_objects_on_conduct):  # noqa: F811
    """🔴 Plane №520: приложение следует стадии ОБЪЕКТА, а не наименьшей.

    Объект А закончил ознакомление и идёт на «Проведении», объект Б отстал —
    `event.stage` падает до стадии отстающего (с №412 это НАИМЕНЬШАЯ стадия
    среди объектов). Документ объекта А приложение терять не должен: лист
    подписан именно им.

    Мутация, которую стережёт проба: вернуть в `acknowledgement_completed`
    сравнение по `event.stage` — документ первого объекта останется без
    строки «Приложение. Лист ознакомления», и проба покраснеет.
    """
    _, event_id, first, second = two_objects_on_conduct
    from organization_management.apps.ops import security_events as service

    second.stage = "ACKNOWLEDGEMENT"
    second.save(update_fields=["stage", "updated_at"])
    event = service.lock_event(event_id)
    service.recompute_event_stage(event)
    event.refresh_from_db()
    assert event.stage == "ACKNOWLEDGEMENT", "мероприятие не опустилось до отстающего — проба ничего не стережёт"

    ahead = _text(render_placement(event.code, fmt="docx", visit_object_id=str(first.pk)))
    assert "Приложение. Лист ознакомления" in ahead

    # А у отстающего объекта приложения нет — порог работает в обе стороны.
    behind = _text(render_placement(event.code, fmt="docx", visit_object_id=str(second.pk)))
    assert "Приложение. Лист ознакомления" not in behind


def test_registry_knows_the_case_and_the_endpoint_serves_it(manager, two_objects_on_conduct):  # noqa: F811
    _, event_id, first, _ = two_objects_on_conduct
    from organization_management.apps.ops import security_events as service
    event = service.lock_event(event_id)
    kinds = {k["kind"]: k for k in documents_registry.list_kinds()}
    assert kinds["case"]["needsEvent"] is True
    payload, name = documents_registry.render("case", event_code=event.code, fmt="docx", visit_object_id=str(first.pk))
    assert name == f"delo-{event.code}.docx" and len(payload) > 1000
    resp = manager.get(
        f"/api/ops/event-documents/render/?kind=case&event={event.code}&ext=docx&visitObject={first.pk}"
    )
    assert resp.status_code == 200, resp.content


# ── Оценки в деле охраняются своим правом (Plane №695) ──────────────────


def test_the_case_hides_the_evaluations_from_a_plain_viewer(
    viewer, two_objects_on_conduct  # noqa: F811
):
    """Баллы и комментарии по людям — не для того, у кого только `event.view`.

    🔴 ЧТО БЫЛО НЕ ТАК. Дело вшивает оценку и комментарий по КАЖДОМУ
    сотруднику, а ручка выдачи документа закрыта `event.view`, тогда как ручка
    оценок требует `event.manage`. Роли `EVENT_APPROVER`, `PATROL_LEAD`,
    `DUTY_PLANNER`, `AUDITOR` скачивали то, что им во всех прочих местах
    закрыто. Раздел объявляет своим правилом «выгрузка открывает ровно то, что
    показывают экраны» — здесь оно и нарушалось.
    """
    _, event_id, first, _ = two_objects_on_conduct
    from organization_management.apps.ops import security_events as service

    event = service.lock_event(event_id)
    scored = [
        row
        for row in (event.placement_assignments or [])
        if row.get("employeeName")
    ]
    assert scored, "фикстура обязана дать хоть одного назначенного"

    response = viewer.get(
        f"/api/ops/event-documents/render/?kind=case&event={event.code}"
        f"&ext=docx&visitObject={first.pk}"
    )

    assert response.status_code == 200, response.content
    import base64

    text = _text(base64.b64decode(response.json()["contentBase64"]))
    assert "Оценки сотрудников" in text, (
        "раздел обязан остаться в оглавлении: дело не должно молча менять состав"
    )
    assert "нет права" in text, "изъятие обязано быть названо вслух, а не скрыто"
    # Отрезается ИМЕННО раздел оценок: фамилия того же человека законно стоит
    # в расстановке и в листе ознакомления — они открыты по `event.view`.
    # Закрыты балл и комментарий, а не факт назначения на пост.
    section = text.split("Оценки сотрудников")[1].split("Журнал штаба")[0]
    assert scored[0]["employeeName"] not in section, (
        "фамилии оценённых не должны попадать в раздел оценок"
    )
    assert "Оценено" not in section, "счёт оценённых — тоже сведение об оценках"


def test_the_case_keeps_the_evaluations_for_the_manager(
    manager, two_objects_on_conduct  # noqa: F811
):
    """Кому оценки открыты — тот получает их как прежде.

    Мутация «прятать раздел всегда» краснеет здесь: починка не должна отобрать
    дело у того, кто ведёт мероприятие.
    """
    _, event_id, first, _ = two_objects_on_conduct
    from organization_management.apps.ops import security_events as service

    event = service.lock_event(event_id)

    text = _text(
        documents_case.render_case(
            event.code,
            visit_object_id=str(first.pk),
            fmt="docx",
            permissions={"event.view", "event.manage"},
        )
    )

    assert "Оценено" in text
    assert "нет права" not in text


def test_the_case_without_permissions_keeps_the_old_behaviour(
    manager, two_objects_on_conduct  # noqa: F811
):
    """Прямой вызов без перечня прав собирает дело целиком.

    Сборщик зовут не только ручка: у него нет собственного понятия о том, кто
    спрашивает. Значение по умолчанию `None` означает «прав не проверяем» —
    иначе внутренний вызов молча терял бы раздел, а не отказывал.
    """
    _, event_id, first, _ = two_objects_on_conduct
    from organization_management.apps.ops import security_events as service

    event = service.lock_event(event_id)

    text = _text(
        documents_case.render_case(
            event.code, visit_object_id=str(first.pk), fmt="docx"
        )
    )

    assert "Оценено" in text

# ── Приложение у ОМ БЕЗ объектов посещения (Plane №697) ─────────────────


def test_placement_of_an_event_without_visit_objects_still_gets_the_sheet(
    manager, two_objects_on_conduct  # noqa: F811
):
    """Лист ознакомления печатается и там, где объектов посещения нет.

    🔴 ЧТО БЫЛО НЕ ТАК. Приложение `[ОЗН-07]` стояло под условием
    `visit is not None`, а мероприятия без объектов посещения достижимы —
    `_document_target` для них СОЗНАТЕЛЬНО возвращает `None` и собирает весь
    расчёт мероприятия («сохранение живых данных: у таких ОМ расчёт лежит в
    мероприятии»). В итоге `render_case` печатал лист, а `render_placement`
    для того же ОМ — нет: два документа об одном мероприятии расходились.

    Сборщик строк листа `visit=None` умеет с самого начала (берёт все посты
    расчёта) — не хватало ровно снятия условия.
    """
    _, event_id, _first, _second = two_objects_on_conduct
    from organization_management.apps.ops import security_events as service

    event = service.lock_event(event_id)
    # Объекты СНИМАЮТСЯ намеренно: предмет пробы — ОМ, у которого их нет
    # вовсе. Заводить такой цепочкой с нуля значило бы проверять путь
    # создания, а не сборку документа.
    event.visit_objects.all().delete()

    text = _text(render_placement(event.code, fmt="docx"))

    assert "Приложение. Лист ознакомления" in text
    # Лист не пустой: посты берутся из расчёта мероприятия.
    assigned = [a for a in event.placement_assignments if a.get("employeeName")]
    assert assigned and assigned[0]["employeeName"] in text


def test_the_sheet_is_still_absent_before_acknowledgement_without_objects(
    manager, two_objects_on_conduct  # noqa: F811
):
    """Порог стадии не отменён: без объектов он тот же самый.

    Мутация «печатать приложение всегда» краснеет здесь — иначе починка №697
    заодно снесла бы правило `[ОЗН-07]` «после завершения ознакомления».
    """
    _, event_id, _first, _second = two_objects_on_conduct
    from organization_management.apps.ops import security_events as service

    event = service.lock_event(event_id)
    event.visit_objects.all().delete()
    event.stage = "PLACEMENT"
    event.save(update_fields=["stage", "updated_at"])

    text = _text(render_placement(event.code, fmt="docx"))

    assert "Приложение. Лист ознакомления" not in text


def test_the_case_names_the_post_of_a_remark_instead_of_its_id(
    manager, two_objects_on_approval  # noqa: F811
):
    """Колонка «Привязка» печатает «сектор · пост», а не `post-…` (Plane №865).

    ЧТО СТЕРЕЖЁТСЯ. Подпись бралась из ключа `postName`, которого замечанию не
    пишет НИКТО, поэтому ветка была мертва и в дело уезжал сырой
    идентификатор. Экранная половина при этом была в порядке — расхождение
    видел только тот, кто открывал файл, а проб на содержимое колонки не было
    вовсе.

    🔴 КРАСНОТА НА МУТАЦИИ: верни в `_remark_attachment` строку
    `named = remark.get("postName") or remark.get("postId")` — в тексте дела
    появится идентификатор, и обе проверки ниже покраснеют.
    """
    from organization_management.apps.ops import security_events as service

    _, event_id, first, _, _ = two_objects_on_approval
    event = service.lock_event(event_id)
    post = next(
        p for p in event.recon_sector_posts if str(p.get("visitObjectId")) == str(first.pk)
    )
    first.approval_remarks = [
        {
            "id": "rem-1",
            "text": "Пост просматривается не полностью",
            "authorName": "Согласующий",
            "status": "OPEN",
            "postId": post["id"],
        }
    ]
    first.save(update_fields=["approval_remarks", "updated_at"])

    text = _text(
        documents_case.render_case(
            event.code, visit_object_id=str(first.pk), fmt="docx"
        )
    )

    assert service.post_label(post) in text, "дело не назвало пост замечания"
    assert str(post["id"]) not in text, "в дело уехал идентификатор поста"


def test_the_case_still_names_a_post_taken_off_the_calculation(
    manager, two_objects_on_approval  # noqa: F811
):
    """Отвязанное замечание по-прежнему называет снятый пост (Plane №510).

    Вторая половина той же ячейки, и она проверяется вместе с первой
    намеренно: правка №865 трогает порядок веток, а именно в порядке и живёт
    поведение «снятый пост называется, а не превращается в „общее“».
    """
    _, event_id, first, _, _ = two_objects_on_approval
    from organization_management.apps.ops import security_events as service

    event = service.lock_event(event_id)
    first.approval_remarks = [
        {
            "id": "rem-2",
            "text": "Замечание по снятому посту",
            "status": "OPEN",
            "postId": None,
            "detachedPost": "Сектор А · Пост 7",
        }
    ]
    first.save(update_fields=["approval_remarks", "updated_at"])

    text = _text(
        documents_case.render_case(
            event.code, visit_object_id=str(first.pk), fmt="docx"
        )
    )

    assert "Сектор А · Пост 7 (пост снят с расчёта)" in text


def test_the_attachment_column_names_the_post_and_the_removed_one_too():
    """🔴 КОЛОНКА «ПРИВЯЗКА» НЕ БЫЛА ЗАКРЫТА НИЧЕМ (Plane №510 и №865, найдено
    ревью №825).

    Функция чистая, а пробы у неё не было ни одной: мутации «печатать
    `postId` вместо подписи» (возврат мёртвой ветки №865) и «выкинуть ветку
    отвязанного поста» (возврат №510) переживали весь набор молча. При этом
    речь о документе, который подписывают и рассылают: сырой `post-4f2a…`
    вместо «Сектор 1 · Пост 2» человек читает как сбой бланка, а «общее»
    вместо «пост снят с расчёта» — как неправду о том, что писал согласующий.

    Четыре положения разом, потому что различаются они одной строкой:
    привязанный пост, отвязанный, готовая подпись из `postName` (запас на
    будущее — побеждает расчёт) и честное «общее».
    """
    labels = {"post-1": "Сектор 1 · Пост 2"}
    attach = documents_case._remark_attachment

    assert attach({"postId": "post-1"}, labels) == "Сектор 1 · Пост 2"
    assert attach({"postId": None, "detachedPost": "Сектор 1 · Пост 2"}, labels) == (
        "Сектор 1 · Пост 2 (пост снят с расчёта)"
    )
    assert attach({"postId": None}, labels) == "общее"
    assert attach({"postId": "post-1", "postName": "Своя подпись"}, labels) == (
        "Своя подпись"
    )
    # Пост, которого в расчёте уже нет, а отвязать забыли: идентификатор —
    # последний запас, но он не должен вытеснять подпись у известного поста.
    assert attach({"postId": "post-404"}, labels) == "post-404"
