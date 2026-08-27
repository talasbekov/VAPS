"""Документ «Список броней в ГОН» (Plane №216).

Восьмой образец задачи №156. Пробы стерегут ТРИ вещи: состав колонок взят из
образца, строки — действующие машины реестра, а снятые с эксплуатации в
документ не попадают.
"""
import pytest
from docx import Document

from organization_management.apps.operations.models_vehicle import OpsVehicle
from organization_management.apps.ops.documents_registry import list_kinds, render

pytestmark = pytest.mark.django_db

#: Подписи колонок — ДОСЛОВНО из образца `04 Список броней в ГОН`. Правятся
#: осознанно и только вместе с образцом: документ обязан совпадать с ним
#: колонка в колонку, иначе он перестаёт быть тем документом.
COLUMNS = [
    "№ п/п",
    "Марка автомобиля",
    "Классификация по кузову",
    "Год выпуска",
    "ГРНЗ",
    "Класс брони",
    "Дислокация",
    "Примечание",
]


def _car(**over):
    data = {
        "brand": "Mercedes-Benz S680 Maybach 4 М (брон.)",
        "body_class": "седан (223)",
        "production_year": 2023,
        "plate": "111 aa 01",
        "armor_class": "VR7",
        "deployment": "Астана",
        "note": "Автохозяйство",
    }
    data.update(over)
    return OpsVehicle.objects.create(**data)


def _document(tmp_path):
    payload, name = render("vehicles", fmt="docx")
    path = tmp_path / name
    path.write_bytes(payload)
    return Document(str(path))


def test_the_kind_is_offered_by_the_registry_and_needs_no_event():
    """Вид документа виден экрану и мероприятия не требует.

    «Список броней» — про парк, а не про ОМ: спрашивать код мероприятия у
    перечня машин автохозяйства не за чем.
    """
    kinds = {row["kind"]: row for row in list_kinds()}

    assert kinds["vehicles"]["label"] == "Список броней в ГОН"
    assert kinds["vehicles"]["needsEvent"] is False


def test_the_table_repeats_the_sample_column_for_column(tmp_path):
    """Шапка таблицы совпадает с образцом.

    Красная на мутации: переименуй любую колонку в шаблоне — проба назовёт
    именно её.
    """
    _car()

    table = _document(tmp_path).tables[0]

    header = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
    assert header == COLUMNS


def test_rows_are_the_live_registry_not_an_empty_form(tmp_path):
    """Строки — машины реестра со всеми сведениями образца.

    Это и отличает документ от бланка: он показывает то, что в системе есть,
    а не место, куда надо вписать руками.
    """
    _car()
    _car(plate="222 bb 02", brand="Toyota Land Cruiser 300 (брон.)", armor_class="VR6")

    table = _document(tmp_path).tables[0]

    assert len(table.rows) == 3  # шапка и две машины
    first = [cell.text.strip() for cell in table.rows[1].cells]
    assert first[0] == "1"
    assert first[1] == "Mercedes-Benz S680 Maybach 4 М (брон.)"
    assert first[2] == "седан (223)"
    assert first[3] == "2023"
    assert first[4] == "111 aa 01"
    assert first[5] == "VR7"
    assert first[6] == "Астана"
    assert first[7] == "Автохозяйство"


def test_a_retired_car_is_not_in_the_list(tmp_path):
    """Снятая машина в «что есть в парке» не входит.

    Красная на мутации: убери `is_active=True` из отбора — снятая машина
    появится в документе, и список перестанет отвечать на свой вопрос.
    """
    _car()
    _car(plate="333 cc 03", brand="Списанная машина", is_active=False)

    table = _document(tmp_path).tables[0]

    plates = [row.cells[4].text.strip() for row in table.rows[1:]]
    assert plates == ["111 aa 01"]


def test_a_year_that_is_unknown_prints_empty_and_not_zero(tmp_path):
    """Неизвестный год печатается пусто.

    Ноль в колонке «Год выпуска» читатель принял бы за год, а пустая клетка —
    честный ответ «сведений нет» (то же правило, что у графиков).
    """
    _car(production_year=None)

    table = _document(tmp_path).tables[0]

    assert table.rows[1].cells[3].text.strip() == ""
