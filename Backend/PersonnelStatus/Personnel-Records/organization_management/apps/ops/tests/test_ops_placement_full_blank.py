"""Бланк «Общая расстановка» выгружается тем же файлом (Plane №164).

ЧТО ЗДЕСЬ СТЕРЕЖЁТСЯ И ЧЕГО НЕТ. Решение заказчика — «сделай выгрузку точно
такого же файла, обезлич все внутри» — распадается на три проверяемых
утверждения, и у каждого своя проба:

1. вёрстка образца доехала целиком, а не пересобрана по полям;
2. внутри не осталось ни одного места подстановки — ни видимого `{{…}}`,
   ни забытого человека;
3. даты в бланке — период ЭТОГО мероприятия, а не числа образца.

С 28.08.2026 добавилось четвёртое: люди в бланке стоят ПО РОЛИ (Plane №240).
Заказчик ответил на вопрос №195 вариантом «б» — справочник ролей наряда (№237)
и роль у назначения (№238), — и прежняя проба «людей не подставляют вовсе»
переписана осознанно, как её докстрока и требовала.
"""
import datetime as dt
import io
import re
import zipfile

import pytest

from organization_management.apps.ops import documents_registry as registry
from organization_management.apps.ops import documents_placement_full as blank

pytestmark = pytest.mark.django_db


@pytest.fixture
def event():
    from organization_management.apps.ops import security_events as event_service

    record = event_service.create_event(
        title="Проба бланка расстановки",
        object_id=None,
        business_date=dt.date(2026, 4, 20).isoformat(),
        kind="FOREIGN",
        actor="test",
    )
    record.business_date_end = dt.date(2026, 4, 21)
    record.save(update_fields=["business_date_end"])
    return record


def docx_text(payload):
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        raw = archive.read("word/document.xml").decode("utf-8", "ignore")
    return re.sub(r"<[^>]+>", " ", raw)


def test_the_blank_keeps_the_customers_own_layout(event):
    """Четырнадцать таблиц образца доезжают в выгрузку.

    Это и есть «точно такой же файл»: пересборка по полям дала бы одну
    таблицу — ровно то, что делает соседний вид «Расстановка».
    """
    from docx import Document

    payload, name = registry.render("placement_full", event_code=event.code, fmt="docx")

    assert name.endswith(".docx") and event.code in name
    document = Document(io.BytesIO(payload))
    assert len(document.tables) == 14, (
        f"вёрстка образца потеряна: таблиц {len(document.tables)} вместо 14"
    )


def test_nothing_is_left_unfilled_in_the_blank(event):
    """Ни одного `{{…}}` в выгруженном файле.

    Место подстановки, доехавшее до заказчика, читается как поломка выгрузки,
    а не как «система этого не знает».
    """
    payload, _ = registry.render("placement_full", event_code=event.code, fmt="docx")

    left = blank.PLACEHOLDER.findall(docx_text(payload))

    assert left == [], f"в бланке остались места подстановки: {sorted(set(left))[:10]}"


def test_the_dates_are_the_period_of_this_event(event):
    """Даты бланка — период мероприятия, а не числа образца."""
    payload, _ = registry.render("placement_full", event_code=event.code, fmt="docx")

    text = docx_text(payload)

    assert "20-21.04.2026" in text, "период мероприятия в бланк не попал"
    assert "22.04.2026" not in text, "в бланке остались даты образца"


def test_a_single_day_event_is_not_written_as_a_range(event):
    """Однодневное мероприятие пишется одной датой, а не «20-20».

    Образец пишет и диапазоном, и одной датой; выдумывать второй день ради
    единообразия значило бы сообщить о мероприятии неправду.
    """
    event.business_date_end = None
    event.save(update_fields=["business_date_end"])

    payload, _ = registry.render("placement_full", event_code=event.code, fmt="docx")

    text = docx_text(payload)
    assert "20.04.2026" in text
    assert "20-20.04.2026" not in text


def test_people_stand_at_the_places_of_their_role(event):
    """Человек стоит там, где написана ЕГО роль (Plane №240).

    🔴 Это и есть ответ на находку №195: до неё бланк заполнялся бы порядком
    следования, то есть водителем VIP становился человек с поста оцепления.
    Проба заводит ДВЕ роли и двух людей крест-накрест и требует, чтобы каждый
    оказался у своей подписи, — совпадение по порядку тут ничего не докажет.
    """
    from organization_management.apps.operations.models import OpsDictionaryEntry
    from organization_management.apps.ops import documents_placement_full as blank

    OpsDictionaryEntry.objects.create(
        dictionary_code="PLACEMENT_ROLES", code="DRIVER_VIP",
        label="Водитель VIP (VIP жүргізушісі)", is_active=True,
    )
    OpsDictionaryEntry.objects.create(
        dictionary_code="PLACEMENT_ROLES", code="MOTORCADE_LEAD",
        label="Ответственный за кортеж (Кортежге жауапты)", is_active=True,
    )
    # 🔴 ПРОБА ДОПОЛНЕНА СЕКЦИЕЙ ОСОЗНАННО (Plane №242, Ш-6). До неё людям
    # хватало роли: раскладка шла по ней одной. Теперь ключ места — ПАРА
    # (роль, секция), и назначение без секции в место с секцией не встаёт
    # намеренно. Секция берётся у САМОГО МЕСТА, а не выдумывается: проба
    # проверяет крест-накрест по ролям, и подставлять чужую секцию значило бы
    # менять её предмет.
    roles = blank.placeholder_roles()
    sections = blank.placeholder_sections()
    vip_places = [name for name, code in roles.items() if code == "DRIVER_VIP"]
    motorcade_places = [name for name, code in roles.items() if code == "MOTORCADE_LEAD"]
    assert vip_places and motorcade_places, "в бланке не нашлось мест этих ролей — проба вакуумна"

    event.placement_assignments = [
        {"id": "a-1", "postId": "p1", "employeeId": "1",
         "employeeName": "Кортежев К.", "roleCode": "MOTORCADE_LEAD",
         "sectionCode": sections.get(motorcade_places[0])},
        {"id": "a-2", "postId": "p2", "employeeId": "2",
         "employeeName": "Випов В.", "roleCode": "DRIVER_VIP",
         "sectionCode": sections.get(vip_places[0])},
    ]
    event.save(update_fields=["placement_assignments"])

    values = blank.placement_full_values(event)

    assert values[vip_places[0]] == "Випов В."
    assert values[motorcade_places[0]] == "Кортежев К."
    # Крест-накрест: каждый стоит ТОЛЬКО у своей подписи.
    assert values[motorcade_places[0]] != "Випов В."


def test_places_without_a_role_of_their_own_stay_empty(event):
    """Перечисления («роль: X, Y, Z») не заполняются — какой из них чей,
    система не знает, и догадка здесь была бы тем же «наугад».
    """
    from organization_management.apps.ops import documents_placement_full as blank

    values = blank.placement_full_values(event)
    roles = blank.placeholder_roles()

    unlabelled = [
        name for name in values
        if name.startswith("person_") and name not in roles
    ]
    assert unlabelled, "в бланке не осталось мест без подписи — образец подменили"
    assert {values[name] for name in unlabelled} == {""}


def test_an_unknown_event_is_refused_by_code(event):
    """Несуществующий код — отказ 404, а не пустой бланк."""
    from organization_management.apps.operations.exceptions import DomainError

    with pytest.raises(DomainError) as raised:
        registry.render("placement_full", event_code="НЕТ-ТАКОГО", fmt="docx")

    assert raised.value.http_status == 404


# ── Ш-1 плана №242: секция места читается из бланка ──────────────────────────


def test_sections_are_read_from_the_template_not_guessed():
    """Каждое место знает СВОЮ секцию, а не первую в документе.

    Красная проба к Plane №242. Одной роли месту не хватает: «Көшпелі
    күзетінің жауаптысы» есть у восьми выездных охран подряд, и раскладка по
    роли ставила первого назначенного в первую охрану наугад.

    Проверяются РАЗНЫЕ секции у мест из разных разделов — мутация «брать
    первый заголовок документа» проходит проверку «секция есть» и краснит
    именно здесь.
    """
    sections = blank.placeholder_sections()

    assert sections["person_2"] == "ULAN_BATOR_KOSHPELI_KUZET"
    assert sections["person_19"] == "TASHKENT_KOSHPELI_KUZET"
    assert sections["person_34"] == "EREVAN_KOSHPELI_KUZET"
    assert (
        sections["person_2"] != sections["person_19"] != sections["person_34"]
    ), "секции разных выездных охран слились в одну"


def test_a_place_above_the_first_heading_has_no_section():
    """Место ВЫШЕ первого заголовка секции не получает её выдуманной.

    `person_1` — ответственный за ВСЮ выездную охрану, он стоит до первого
    раздела. Приписать ему «Ұлан-батор» значило бы вернуть ту самую догадку,
    от которой уходили в №195: документ назвал бы человека ответственным за
    один кортеж вместо всех.
    """
    assert "person_1" not in blank.placeholder_sections()


def test_every_other_place_belongs_to_a_section():
    """Правило заголовка накрывает ВЕСЬ бланк, а не удобную его часть.

    Числа замерены на образце и стоят здесь нарочно: если следующий образец
    заказчика придёт с разделом другой формы, проба скажет об этом сразу, а не
    оставит десяток мест молча пустыми.
    """
    places = set(blank.template_placeholders())
    people = {name for name in places if name.startswith("person_")}
    with_section = set(blank.placeholder_sections())

    assert len(people) == 1027
    assert people - with_section == {"person_1"}


def test_sections_are_deduplicated_by_code():
    """Одна подпись — одна секция, сколько бы раз она ни встретилась.

    «Сапар» объектісі стоит в файле трижды: мероприятие заходит на объект по
    разу на каждый день. Это ОДНА секция, и в справочник (Ш-2) она обязана
    попасть один раз — иначе человек выбирал бы из трёх одинаковых строк.
    """
    sections = blank.template_sections()
    codes = [entry["code"] for entry in sections]

    assert len(codes) == len(set(codes)), "в списке секций есть дубли по коду"
    assert "SAPAR_OBEKTISI" in codes or any("SAPAR" in code for code in codes)


def test_section_code_does_not_depend_on_the_order_in_the_file():
    """Код секции выводится из ПОДПИСИ, а не из её номера в документе.

    Бланк переснимается при каждом новом образце заказчика. Код, зависящий от
    порядка, при вставке раздела в середину молча переехал бы на чужую секцию —
    и назначения людей, сделанные до пересъёмки, указали бы не туда.
    """
    first = blank._section_code("«Ұлан-батор» көшпелі күзет")
    same = blank._section_code("«Ұлан-батор» көшпелі күзет")
    other = blank._section_code("«Ташкент» көшпелі күзет")

    assert first == same
    assert first != other


def test_long_prose_with_quotes_is_not_a_heading():
    """Кавычки в прозе не делают абзац заголовком.

    В бланке есть абзацы вроде «ҚК ӘП (жауынгерлік топ) – 5 қызм. … „Алмаз-1“,
    ІІМ …» — длинные перечисления сил, где кавычки стоят как обычная
    пунктуация. Прими их за заголовки — и половина мест уехала бы в секции,
    которых не существует.
    """
    prose = (
        "ҚК ӘП (жауынгерлік топ) – 5 қызм. Арлан тәулікте, ҚМ ҰҰА қарсы "
        "іс-қимылдың құрылғысы 1 бірлік «Алмаз-1», ІІМ АМБ мерген – 1 жұп"
    )
    lead_in = "Жедел жасақтың іс-қимылдарын үйлестіруді жедел штаб «Балхаш» жүзеге асырады:"

    assert not blank._is_section_heading(prose)
    assert not blank._is_section_heading(lead_in)
    assert blank._is_section_heading("«Ұлан-батор» көшпелі күзет")


# ── Ш-6 плана №242: раскладка по ПАРЕ (роль, секция) ─────────────────────────


def test_one_role_in_two_sections_goes_to_its_own_places(event):
    """Два человека с ОДНОЙ ролью и РАЗНЫМИ секциями встают каждый в своё место.

    🔴 Это и есть ответ на №242. «Көшпелі күзетінің жауаптысы» есть у восьми
    выездных охран подряд, и раскладка по одной роли ставила первого
    назначенного в первую охрану — то есть наугад, ровно как до №240, только
    незаметнее: документ выглядел заполненным.

    Проба берёт места ОДНОЙ роли в РАЗНЫХ секциях и требует, чтобы каждый
    человек оказался в своей. Совпадение по порядку здесь ничего не докажет:
    людей ставят крест-накрест — первым идёт тот, чья секция вторая.
    """
    from organization_management.apps.operations.models import OpsDictionaryEntry
    from organization_management.apps.ops import documents_placement_full as blank

    # Роль взята та, у которой места есть в НЕСКОЛЬКИХ секциях бланка
    # («Кортежге жауапты» — в каждой из восьми выездных охран). Роль из одной
    # секции сделала бы пробу вакуумной: сравнивать было бы не с чем.
    role_code = "MOTORCADE_LEAD"
    OpsDictionaryEntry.objects.create(
        dictionary_code="PLACEMENT_ROLES", code=role_code,
        label="Ответственный за кортеж (Кортежге жауапты)",
        is_active=True,
    )
    # Роли читаются ПОСЛЕ заведения записи: `placeholder_roles` сверяет
    # подписи бланка со СПРАВОЧНИКОМ, и на пустом справочнике мест не находит
    # вовсе (первая редакция пробы упала именно так и была вакуумна).
    roles = blank.placeholder_roles()
    sections = blank.placeholder_sections()
    places = [name for name, code in roles.items() if code == role_code]
    by_section = {}
    for name in places:
        section = sections.get(name)
        if section is not None and section not in by_section:
            by_section[section] = name
    assert len(by_section) >= 2, (
        "в бланке нет одной роли в двух секциях — проверять №242 не на чем"
    )
    (first_section, first_place), (second_section, second_place) = list(
        by_section.items()
    )[:2]

    # КРЕСТ-НАКРЕСТ: первым в списке идёт человек ВТОРОЙ секции.
    event.placement_assignments = [
        {"id": "s-1", "postId": "p1", "employeeId": "1",
         "employeeName": "Второв В.", "roleCode": role_code,
         "sectionCode": second_section},
        {"id": "s-2", "postId": "p2", "employeeId": "2",
         "employeeName": "Первов П.", "roleCode": role_code,
         "sectionCode": first_section},
    ]
    event.save(update_fields=["placement_assignments"])

    values = blank.placement_full_values(event)

    assert values[first_place] == "Первов П."
    assert values[second_place] == "Второв В."


def test_an_assignment_without_a_section_does_not_wander_into_sections(event):
    """Назначение БЕЗ секции не расползается по разделам бланка.

    Расстановки, сделанные до №242, секции не несут. Поставь их в первое
    попавшееся место роли — и документ снова назовёт человека ответственным за
    чужой кортеж, только теперь молча и задним числом. Пустое место честнее:
    оно означает «система этого не знает», как и всегда.
    """
    from organization_management.apps.operations.models import OpsDictionaryEntry
    from organization_management.apps.ops import documents_placement_full as blank

    # Роль взята та, у которой места есть в НЕСКОЛЬКИХ секциях бланка
    # («Кортежге жауапты» — в каждой из восьми выездных охран). Роль из одной
    # секции сделала бы пробу вакуумной: сравнивать было бы не с чем.
    role_code = "MOTORCADE_LEAD"
    OpsDictionaryEntry.objects.create(
        dictionary_code="PLACEMENT_ROLES", code=role_code,
        label="Ответственный за кортеж (Кортежге жауапты)",
        is_active=True,
    )
    roles = blank.placeholder_roles()
    sections = blank.placeholder_sections()
    with_section = [
        name for name, code in roles.items()
        if code == role_code and sections.get(name) is not None
    ]
    assert with_section, "мест этой роли с секцией в бланке нет — проба вакуумна"

    event.placement_assignments = [
        {"id": "s-3", "postId": "p1", "employeeId": "1",
         "employeeName": "Безсекциев Б.", "roleCode": role_code},
    ]
    event.save(update_fields=["placement_assignments"])

    values = blank.placement_full_values(event)

    assert {values[name] for name in with_section} == {""}, (
        "человек без секции встал в место с секцией — раскладка снова гадает"
    )
