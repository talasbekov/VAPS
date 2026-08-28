"""Рендерер .docx расхода (порт apps/documents/generators/expense_docx.py из
Backend/VAPS).

Третья форма одной выгрузки, и единственная ПЕЧАТНАЯ: .csv отдаёт числа для
дальнейшего счёта, .xlsx — лист для работы, а .docx — страницу, которую
подписывают и подшивают. Отсюда всё её своеобразие: A4 landscape, титул 16pt,
таблица 12pt, а поимённый состав в ячейке — 8pt отдельными абзацами.

ДВА РАЗМЕРА В ЯЧЕЙКЕ — то, ради чего формат и портирован. У .xlsx шрифт один
на ячейку, и там состав приходится вписывать в тот же кегль, что и число;
здесь число видно как число, а два десятка фамилий под ним не спорят с ним за
внимание и умещаются на листе.

Times New Roman ставится ЯВНО на каждом run. Дефолт шаблона python-docx —
Calibri, и молчаливое согласие с ним дало бы документ, не похожий на все
остальные бумаги раздела; `w:cs` добивается через rFonts, иначе кириллица и
казахские глифы у части читателей уезжают в подстановочный шрифт.

Рендерер НИЧЕГО не считает и не пересортировывает: числа, порядок строк и
порядок колонок приходят готовыми от билдера и справочника. python-docx
импортируется ЛЕНИВО — раздел не должен требовать библиотеку выгрузки ради
сдачи дня и чтения расхода.
"""
from organization_management.apps.operations.expense_layout import (
    event_cell,
    FIXED_HEAD,
    FONT_NAME,
    MEMBER_SIZE_PT,
    TABLE_SIZE_PT,
    TITLE_SIZE_PT,
    TOTALS_LABEL,
    document_title,
    header_row,
    member_lines,
)


def generate_expense_docx(data):
    """Данные документа → байты .docx."""
    import io

    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt

    def styled(paragraph, text, size_pt, bold=False):
        """Run с ЯВНЫМИ именем и размером шрифта (см. модульную преамбулу)."""
        run = paragraph.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(size_pt)
        run.bold = bold
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), FONT_NAME)
        return run

    def fill_status_cell(cell, cell_data):
        """Число 12pt, под ним состав 8pt — каждый человек своим абзацем.

        Абзацами, а не переводами строки внутри одного run: у абзаца свой
        размер, а перенос внутри run унаследовал бы кегль числа — и вся
        разница с .xlsx пропала бы.
        """
        styled(cell.paragraphs[0], str(cell_data.count), TABLE_SIZE_PT)
        for line in member_lines(cell_data.members):
            styled(cell.add_paragraph(), line, MEMBER_SIZE_PT)

    document = Document()

    # A4 landscape. Дефолт шаблона — Letter portrait, а смена ориентации
    # габариты сама НЕ свопает: ставим и то, и другое, иначе таблица на
    # полтора десятка колонок уходит за край страницы.
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)

    styled(document.add_paragraph(), document_title(data), TITLE_SIZE_PT, bold=True)

    head = header_row(data.columns)
    table = document.add_table(rows=1, cols=len(head))
    # Сетка нужна именно этой форме: в ячейках стоят многострочные списки, и
    # без границ на печати состав одной колонки читается продолжением
    # соседней.
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, head):
        styled(cell.paragraphs[0], label, TABLE_SIZE_PT, bold=True)

    for number, row in enumerate(data.rows, start=1):
        cells = table.add_row().cells
        for index, value in enumerate(
            (number, row.name, row.staff_total, row.list_total, row.vacancies)
        ):
            styled(cells[index].paragraphs[0], str(value), TABLE_SIZE_PT)
        for offset, column in enumerate(data.columns, start=len(FIXED_HEAD)):
            fill_status_cell(cells[offset], row.cells[column])
        # Приданные — «+N», прибавка СВЕРХ списка, и члены у них не бывают
        # (чужие люди). Отличие от источника, где «+N» вешалось на статусную
        # колонку ATTACHED: там прикомандированный внутри списка и приданный
        # сверх него печатались одинаково, хотя складываются по-разному.
        styled(cells[-2].paragraphs[0], f"+{row.attached.count}", TABLE_SIZE_PT)
        styled(cells[-1].paragraphs[0], event_cell(row.event), TABLE_SIZE_PT)

    totals = data.totals
    cells = table.add_row().cells
    # Номера у итога нет: пустая ячейка «№» отделяет его от данных.
    for index, value in enumerate(
        (TOTALS_LABEL, totals.staff_total, totals.list_total, totals.vacancies),
        start=1,
    ):
        styled(cells[index].paragraphs[0], str(value), TABLE_SIZE_PT, bold=True)
    for offset, column in enumerate(data.columns, start=len(FIXED_HEAD)):
        styled(
            cells[offset].paragraphs[0], str(totals.columns[column]), TABLE_SIZE_PT,
            bold=True,
        )
    styled(cells[-2].paragraphs[0], f"+{totals.attached}", TABLE_SIZE_PT, bold=True)
    styled(cells[-1].paragraphs[0], event_cell(totals.event), TABLE_SIZE_PT, bold=True)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
