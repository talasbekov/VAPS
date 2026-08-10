"""Все формы расхода перечисляют людей ОДИНАКОВО.

Формы собирают разные библиотеки: .xlsx — openpyxl, .docx — python-docx, .pdf —
reportlab. Каждая берёт готовые данные и раскладывает их по-своему, и разойтись
они могут незаметно: три отдельных теста, каждый зелёный на своей форме, такого
расхождения не увидят вовсе. А расхождение это не косметика — один и тот же день,
выгруженный дважды в разных форматах, читался бы как два разных документа.

.csv здесь СРАВНИВАЕТСЯ ПО-ДРУГОМУ, и это его договор: машиночитаемая форма несёт
только числа, поимённого состава в ней нет. Проверяется, что его там и правда
нет, — иначе «формы совпадают» однажды окажется правдой по недосмотру.
"""
import io
from datetime import date

import pytest
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from organization_management.apps.operations.expense_csv import generate_expense_csv
from organization_management.apps.operations.expense_docx import generate_expense_docx
from organization_management.apps.operations.expense_document import (
    build_expense_document,
)
from organization_management.apps.operations.expense_pdf import generate_expense_pdf
from organization_management.apps.operations.expense_xlsx import generate_expense_xlsx
from organization_management.apps.operations.strength_report import StatusCatalog

DAY = date(2026, 8, 4)

# Фамилии и уровни подобраны так, чтобы КАЖДЫЙ ключ канона был задействован и ни
# один порядок-самозванец не совпал с ожидаемым: по должности — не по алфавиту,
# по алфавиту — не по номеру строки, и есть человек без должности.
PEOPLE = [
    (1, "Ёлкин", 50),
    (2, "Абрамов", 90),
    (3, "Яковлев", 10),
    (4, "Дроздов", 50),
    (5, "Безродный", None),
]
# Яковлев (10) → Дроздов и Ёлкин (оба 50, по алфавиту: «д» раньше «е», а «ё»
# сравнивается как «е») → Абрамов (90) → Безродный (должности нет — в конец).
EXPECTED = ["Яковлев", "Дроздов", "Ёлкин", "Абрамов", "Безродный"]


@pytest.fixture
def catalog():
    return StatusCatalog.from_rows(
        [
            {
                "code": "IN_SERVICE",
                "priority": 999,
                "report_column_code": "IN_SERVICE",
                "counts_in_staff": True,
            },
            {
                "code": "DUTY",
                "priority": 10,
                "report_column_code": "ON_DUTY",
                "counts_in_staff": True,
            },
        ]
    )


@pytest.fixture
def data(catalog):
    roster = [
        {
            "employee_id": employee_id,
            "full_name": name,
            "rank": "капитан",
            "position_level": level,
        }
        for employee_id, name, level in PEOPLE
    ]
    rows = [
        {
            "employee_id": employee_id,
            "status_type_code": "DUTY",
            "date_start": DAY.isoformat(),
            "date_end": date(2026, 8, 10).isoformat(),
            "source": "USER",
        }
        for employee_id, _, _ in PEOPLE
    ]
    return build_expense_document(
        {"roster": roster, "rows": rows},
        DAY,
        catalog=catalog,
        division_title="Управление",
        staff_total=10,
        vacancies=0,
        attached=0,
    )


def _order_in(text):
    """Порядок наших фамилий в произвольном тексте документа.

    Ищутся ИМЕННО ожидаемые фамилии, а прочий текст игнорируется: у каждой формы
    свои заголовки и подписи, и сравнивать документы целиком значило бы сравнивать
    оформление, а не порядок людей.
    """
    found = [(text.index(name), name) for name in EXPECTED if name in text]
    return [name for _, name in sorted(found)]


def order_from_xlsx(data):
    sheet = load_workbook(io.BytesIO(generate_expense_xlsx(data))).active
    text = "\n".join(
        str(cell.value) for row in sheet.iter_rows() for cell in row if cell.value
    )
    return _order_in(text)


def order_from_docx(data):
    (table,) = DocxDocument(io.BytesIO(generate_expense_docx(data))).tables
    text = "\n".join(
        paragraph.text
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    return _order_in(text)


def order_from_pdf(data):
    reader = PdfReader(io.BytesIO(generate_expense_pdf(data)))
    return _order_in("\n".join(page.extract_text() or "" for page in reader.pages))


# ── Каждая форма — по канону ─────────────────────────────────────────────


def test_the_fixture_actually_distinguishes_the_canon(data):
    """Опора всего файла: ожидаемый порядок не совпадает ни с порядком строк в
    базе, ни с чистым алфавитом.

    Совпади он хоть с одним — все проверки ниже зеленели бы при сортировке по
    чему угодно.
    """
    by_row_id = [name for _, name, _ in PEOPLE]
    alphabetical = sorted(by_row_id)

    assert EXPECTED != by_row_id
    assert EXPECTED != alphabetical


def test_the_sheet_lists_people_by_the_canon(data):
    assert order_from_xlsx(data) == EXPECTED


def test_the_printed_page_lists_people_by_the_canon(data):
    assert order_from_docx(data) == EXPECTED


def test_the_pdf_lists_people_by_the_canon(data):
    assert order_from_pdf(data) == EXPECTED


# ── И одинаково между собой ──────────────────────────────────────────────


def test_all_document_forms_agree_with_each_other(data):
    """Несущий тест: три библиотеки, один порядок.

    Три отдельных теста выше зелены каждый на своей форме — расхождение МЕЖДУ
    ними они бы не увидели.
    """
    assert order_from_xlsx(data) == order_from_docx(data) == order_from_pdf(data)


def test_every_person_reaches_every_form(data):
    """Совпадение порядка ничего не стоит, если кого-то потеряли все трое
    одинаково."""
    assert len(order_from_xlsx(data)) == len(PEOPLE)
    assert len(order_from_docx(data)) == len(PEOPLE)
    assert len(order_from_pdf(data)) == len(PEOPLE)


# ── Машиночитаемая форма — по своему договору ────────────────────────────


def test_the_csv_carries_numbers_and_no_names(data):
    """Договор .csv: только числа. Таблица, в ячейке которой лежит двадцать
    фамилий, не открывается ни одним потребителем CSV так, как он ожидает.

    Проверяется здесь, а не «забыто»: иначе однажды состав протечёт и в неё, а
    тест «формы совпадают» этого не заметит — он про три другие формы.
    """
    text = generate_expense_csv(data).decode("utf-8-sig")

    assert "Управление" in text
    for _, name, _ in PEOPLE:
        assert name not in text
