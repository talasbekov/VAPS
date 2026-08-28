"""Выгрузка .docx: печатная форма расхода.

Проверяется ОТКРЫТЫЙ документ, а не байты: важен не zip внутри (его держит
python-docx), а то, что увидит и подпишет читатель распечатки — кегли, сетка,
ориентация листа и поимённый состав под числом.
"""
import io
from datetime import date

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt

from organization_management.apps.operations.expense_docx import generate_expense_docx
from organization_management.apps.operations.expense_document import (
    build_expense_document,
    combine_documents,
)
from organization_management.apps.operations.expense_layout import (
    CELL_MAX_MEMBERS,
    COLUMN_LABELS,
    FONT_NAME,
    MEMBER_SIZE_PT,
    TABLE_SIZE_PT,
    TITLE_SIZE_PT,
    TOTALS_LABEL,
)
from organization_management.apps.operations.strength_report import StatusCatalog

DAY = date(2026, 8, 4)


@pytest.fixture
def catalog():
    return StatusCatalog.from_rows(
        [
            {
                "code": "IN_SERVICE",
                "priority": 999,
                "report_column_code": "IN_SERVICE",
                "counts_in_staff": True,
            },
            {
                "code": "DUTY",
                "priority": 10,
                "report_column_code": "ON_DUTY",
                "counts_in_staff": True,
            },
            {
                "code": "VACATION",
                "priority": 30,
                "report_column_code": "VACATION",
                "counts_in_staff": True,
            },
        ]
    )


def member(employee_id, full_name="Иванов Иван", rank="капитан"):
    return {"employee_id": employee_id, "full_name": full_name, "rank": rank}


def fact(employee_id, code):
    return {
        "employee_id": employee_id,
        "status_type_code": code,
        "date_start": DAY.isoformat(),
        "date_end": date(2026, 8, 10).isoformat(),
        "source": "USER",
    }


def document(catalog, roster, rows, **overrides):
    kwargs = {
        "catalog": catalog,
        "division_title": "Управление кадров",
        "staff_total": 10,
        "vacancies": 7,
        "attached": 2,
    }
    kwargs.update(overrides)
    return build_expense_document({"roster": roster, "rows": rows}, DAY, **kwargs)


@pytest.fixture
def data(catalog):
    return document(
        catalog,
        [member(1), member(2, full_name="Петров Пётр", rank=""), member(3)],
        [fact(1, "DUTY"), fact(2, "VACATION")],
    )


def opened(data):
    return Document(io.BytesIO(generate_expense_docx(data)))


def table_of(data):
    (table,) = opened(data).tables
    return table


def header_labels(table):
    return [cell.text for cell in table.rows[0].cells]


def column_index(table, label):
    labels = header_labels(table)
    if label not in labels:
        raise AssertionError(f"колонка {label!r} не найдена в шапке")
    return labels.index(label)


def cell_paragraphs(table, row_index, label):
    cell = table.rows[row_index].cells[column_index(table, label)]
    return [paragraph.text for paragraph in cell.paragraphs]


# ── Лист, на котором это печатают ────────────────────────────────────────


def test_the_page_is_a4_landscape(data):
    """Смена ориентации габариты сама не свопает — иначе таблица за краем."""
    (section,) = opened(data).sections

    assert section.orientation == WD_ORIENT.LANDSCAPE
    # Габариты сверяются в миллиметрах: формат хранит их в твипах, и
    # round-trip через них теряет доли — на бумаге незаметные, для сравнения
    # с Mm(297) фатальные.
    assert (round(section.page_width.mm), round(section.page_height.mm)) == (297, 210)


def test_the_title_names_the_division_and_the_date(data):
    (title,) = [
        paragraph
        for paragraph in opened(data).paragraphs
        if paragraph.text.strip()
    ]

    assert "Управление кадров" in title.text
    assert "04.08.2026" in title.text
    (run,) = title.runs
    assert (run.font.size, run.bold) == (Pt(TITLE_SIZE_PT), True)


def test_every_run_carries_the_document_font(data):
    """Дефолт шаблона python-docx — Calibri.

    Молчаливое согласие с ним дало бы документ, не похожий на остальные
    бумаги раздела.
    """
    opened_document = opened(data)
    runs = [run for paragraph in opened_document.paragraphs for run in paragraph.runs]
    for table in opened_document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    runs.extend(paragraph.runs)

    assert runs
    assert {run.font.name for run in runs} == {FONT_NAME}


def test_the_table_is_ruled(data):
    """В ячейках многострочные списки: без границ состав одной колонки
    читается продолжением соседней."""
    assert table_of(data).style.name == "Table Grid"


# ── Шапка и знаменатель ──────────────────────────────────────────────────


def test_the_header_follows_the_catalog(data):
    labels = header_labels(table_of(data))

    assert labels[:5] == ["№", "Управление", "По штату", "По списку", "Вакансии"]
    assert labels[5:] == [
        COLUMN_LABELS["ON_DUTY"],
        COLUMN_LABELS["VACATION"],
        COLUMN_LABELS["IN_SERVICE"],
        "Придано (+N)",
        # «На ОМ (гр./нар.)» — справочная колонка занятости мероприятиями
        # (Plane №243, решение заказчика «просто то что есть печатать»).
        # Стоит ПОСЛЕ «+N», в самом конце: эти люди уже посчитаны в «В строю»,
        # и место среди колонок расхода предлагало бы сложить их со всеми
        # остальными. Пин правится вместе с шапкой и только осознанно.
        "На ОМ (гр./нар.)",
    ]


def test_the_row_carries_the_denominator(data):
    cells = table_of(data).rows[1].cells

    assert [cell.text for cell in cells[:5]] == ["1", "Управление кадров", "10", "3", "7"]


def test_the_totals_row_is_bold_and_has_no_number(data):
    table = table_of(data)
    cells = table.rows[-1].cells

    assert cells[0].text == ""
    assert cells[1].text == TOTALS_LABEL
    assert [cell.text for cell in cells[2:5]] == ["10", "3", "7"]
    assert all(
        run.bold
        for cell in cells[1:]
        for paragraph in cell.paragraphs
        for run in paragraph.runs
    )


# ── Два размера в одной ячейке — то, ради чего формат и нужен ────────────


def test_the_count_and_the_members_are_set_in_different_sizes(data):
    """У .xlsx шрифт один на ячейку; здесь число видно как число."""
    table = table_of(data)
    cell = table.rows[1].cells[column_index(table, COLUMN_LABELS["ON_DUTY"])]
    count, name = cell.paragraphs

    assert (count.text, count.runs[0].font.size) == ("1", Pt(TABLE_SIZE_PT))
    assert name.text == "капитан Иванов Иван — 04.08.2026–10.08.2026"
    assert name.runs[0].font.size == Pt(MEMBER_SIZE_PT)


def test_a_member_without_a_rank_has_no_leading_space(data):
    _, name = cell_paragraphs(table_of(data), 1, COLUMN_LABELS["VACATION"])

    assert name.startswith("Петров Пётр —")


def test_a_cell_without_members_is_just_the_count(data):
    """«В строю» — отсутствие фактов: число есть, брать состав неоткуда."""
    assert cell_paragraphs(table_of(data), 1, COLUMN_LABELS["IN_SERVICE"]) == ["1"]


def test_a_long_list_is_truncated_with_an_honest_tail(catalog):
    """Молча обрезанный список выглядел бы как полный."""
    roster = [member(index, full_name=f"Сотрудник {index}") for index in range(30)]
    rows = [fact(index, "DUTY") for index in range(30)]
    table = table_of(document(catalog, roster, rows))
    paragraphs = cell_paragraphs(table, 1, COLUMN_LABELS["ON_DUTY"])

    assert paragraphs[0] == "30"
    assert len(paragraphs) == CELL_MAX_MEMBERS + 2  # число + имена + хвост
    assert paragraphs[-1] == f"… ещё {30 - CELL_MAX_MEMBERS}"


# ── Приданные ────────────────────────────────────────────────────────────


def test_attached_is_shown_as_a_surplus(data):
    """«+N» — прибавка СВЕРХ списка, и в строке, и в итоге."""
    table = table_of(data)

    assert cell_paragraphs(table, 1, "Придано (+N)") == ["+2"]
    assert cell_paragraphs(table, -1, "Придано (+N)") == ["+2"]


def test_the_attached_status_column_is_a_plain_number(catalog):
    """Прикомандированный внутри списка — не приданный сверх него.

    Отличие от источника, где «+N» вешалось на статусную колонку ATTACHED:
    два числа складываются по-разному, и печататься одинаково не должны.
    """
    catalog = StatusCatalog.from_rows(
        [
            {
                "code": "IN_SERVICE",
                "priority": 999,
                "report_column_code": "IN_SERVICE",
                "counts_in_staff": True,
            },
            {
                "code": "ATTACHED",
                "priority": 10,
                "report_column_code": "ATTACHED",
                "counts_in_staff": True,
            },
        ]
    )
    table = table_of(document(catalog, [member(1)], [fact(1, "ATTACHED")]))

    assert cell_paragraphs(table, 1, COLUMN_LABELS["ATTACHED"])[0] == "1"


# ── Пустая сводка ────────────────────────────────────────────────────────


def test_a_document_without_rows_still_has_a_header_and_totals(catalog):
    """Ни одной строки — законное состояние сводки, но бумага обязана быть.

    Шапку задают колонки справочника, а не первая строка: пустая сводка
    выходит с той же шапкой, что и полная, иначе два дня одного
    подразделения перестали бы сличаться глазом.
    """
    empty = combine_documents(
        [],
        division_title="Департамент",
        business_date=DAY,
        columns=catalog.columns_in_order(),
    )
    table = table_of(empty)

    assert len(table.rows) == 2  # шапка + итог
    assert header_labels(table)[5:] == [
        COLUMN_LABELS["ON_DUTY"],
        COLUMN_LABELS["VACATION"],
        COLUMN_LABELS["IN_SERVICE"],
        "Придано (+N)",
        # «На ОМ (гр./нар.)» — справочная колонка занятости мероприятиями
        # (Plane №243, решение заказчика «просто то что есть печатать»).
        # Стоит ПОСЛЕ «+N», в самом конце: эти люди уже посчитаны в «В строю»,
        # и место среди колонок расхода предлагало бы сложить их со всеми
        # остальными. Пин правится вместе с шапкой и только осознанно.
        "На ОМ (гр./нар.)",
    ]
    assert cell_paragraphs(table, -1, "Придано (+N)") == ["+0"]
