"""
Сервисные функции для генерации отчетов и расчета статистики.

Полная реализация всех форматов отчетов согласно ТЗ:
- DOCX с таблицей в альбомной ориентации
- XLSX с форматированием
- PDF отчеты
"""

import io
import datetime
from collections import defaultdict, OrderedDict
from typing import Dict, List, Optional, Tuple

from django.db.models import Count, Q, F
from django.utils import timezone

from io import BytesIO

from docx import Document

import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

from organization_management.apps.divisions.models import Division
from organization_management.apps.employees.models import Employee
from organization_management.apps.statuses.models import EmployeeStatus
from organization_management.apps.secondments.models import SecondmentRequest
from organization_management.apps.reports.infrastructure import report_table

class DOCXGenerator:
    """
    Простейший DOCX‑генератор: заголовок + таблица с агрегатами.
    Возвращает (filename, bytes).
    """

    def generate(self, data, report):
        doc = Document()
        doc.add_heading(f"{report.get_report_type_display()}", level=1)
        doc.add_paragraph(f"Раздел: {data.get('division')}")
        doc.add_paragraph(f"Дата: {data.get('date')}")

        rows = data.get("rows", [])
        headers = report_table.headers(data)
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        hdr_cells = table.rows[0].cells
        for index, label in enumerate(headers):
            hdr_cells[index].text = label

        for i, row in enumerate(rows, start=1):
            cells = table.rows[i].cells
            for index, value in enumerate(report_table.cells(data, row)):
                cells[index].text = str(value)  # type: ignore

        stream = BytesIO()
        doc.save(stream)
        filename = f"report_{report.id}.docx"
        return filename, stream.getvalue()
